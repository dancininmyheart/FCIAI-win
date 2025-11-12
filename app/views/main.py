import time
import json

import pytz

from flask import Blueprint, render_template, redirect, url_for, flash, request, current_app, send_from_directory, \
    jsonify, session, send_file, abort
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from sqlalchemy.exc import SQLAlchemyError

from config import data_file

# from ..Ingredient_Search.Flask_app import search, download_files
from ..function.adjust_text_size import set_textbox_autofit
from ..function.ppt_translate import process_presentation, process_presentation_add_annotations
from config import base_model_file
from ..models import User, UploadRecord, Translation, StopWord
from ..services.sso_service import get_sso_service
from .. import db
import os
import uuid
import re

from ..utils.task_queue import translation_queue as old_translation_queue
from ..function.ppt_translate_async import process_presentation as process_presentation_async
from ..function.ppt_translate_async import \
    process_presentation_add_annotations as process_presentation_add_annotations_async
from ..utils.enhanced_task_queue import EnhancedTranslationQueue, TranslationTask, translation_queue
from ..utils.thread_pool_executor import thread_pool, TaskType
import openpyxl
from io import BytesIO
import logging
import threading
from datetime import datetime
from app.utils.timezone_helper import format_datetime, datetime_to_isoformat

# from ..utils.Tokenization import Tokenizer
# from ...train import train_model
# sys.stdout.reconfigure(encoding='utf-8')
main = Blueprint("main", __name__)

# 配置日志记录器
logger = logging.getLogger(__name__)

# 使用增强的任务队列替换旧队列
# translation_queue = TranslationQueue()

# 简单任务状态存储（用于公开API）
simple_task_status = {}
simple_task_files = {}

# PDF翻译任务状态存储（用于异步任务）
pdf_task_status_cache = {}
pdf_task_lock = threading.Lock()


def process_pdf_translation_async(pdf_path, original_filename, unique_filename, 
                                  source_lang, target_lang, model, enable_image_ocr,
                                  custom_translations, user_id, task_id):
    """
    异步处理PDF翻译的工作函数
    该函数在独立线程中执行，不阻塞主线程
    """
    from flask import current_app
    from app import create_app
    import zipfile
    import requests
    
    logger = logging.getLogger(__name__)
    logger.info(f"开始异步处理PDF翻译任务: {task_id}")
    
    # 创建应用上下文（异步任务运行在独立线程中，需要显式创建上下文）
    app = create_app()
    with app.app_context():
        try:
            # 获取上传文件夹路径
            project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            upload_folder = current_app.config['UPLOAD_FOLDER']
            if not os.path.isabs(upload_folder):
                upload_folder = os.path.join(project_root, upload_folder)

            pdf_output_dir = os.path.join(upload_folder, 'pdf_outputs')
            os.makedirs(pdf_output_dir, exist_ok=True)
            
            task_work_dir = os.path.join(
                pdf_output_dir,
                f"{os.path.splitext(unique_filename)[0]}_work"
            )
            os.makedirs(task_work_dir, exist_ok=True)

            # 首选方案：使用OSS直链处理PDF
            result = None
            try:
                from app.function.image_ocr.oss_pdf_processor import OSSPDFProcessor
                from app.function.image_ocr.ocr_api import MinerUAPI

                logger.info("初始化OSS PDF处理器")
                logger.info(f"OCR功能状态: {'启用' if enable_image_ocr else '禁用'}")
                oss_processor = OSSPDFProcessor()
                mineru_api = MinerUAPI()

                # 使用OSS直链处理PDF
                logger.info(f"开始使用OSS直链处理PDF: {pdf_path}")
                result = oss_processor.process_pdf_with_mineru(
                    pdf_path, 
                    mineru_api, 
                    bucket="fciai", 
                    region="cn-beijing",
                    enable_ocr=enable_image_ocr
                )

                if result and isinstance(result, dict) and result.get('code') == 0:
                    logger.info("OSS直链方案处理成功")
                else:
                    logger.warning("OSS直链方案处理失败，尝试使用本地PDF处理器...")
                    result = None
            except Exception as e:
                logger.warning(f"OSS直链方案处理失败: {e}")
                result = None

            # 如果OSS直链方案失败，使用本地PDF处理器
            if not result:
                logger.info("使用本地PDF处理器...")
                try:
                    from app.function.local_pdf_processor import LocalPDFProcessor
                    local_processor = LocalPDFProcessor()
                    result = local_processor.process_pdf(pdf_path)
                    logger.info(f"本地PDF处理结果: {result}")
                except Exception as local_e:
                    logger.error(f"本地PDF处理器也失败了: {local_e}")
                    raise Exception('PDF处理失败，请检查文件格式')

            if not result:
                raise Exception("所有PDF处理方法都失败了")

            # 检查结果中的状态码
            if 'code' in result and result['code'] != 0:
                error_msg = result.get('msg', '未知错误')
                raise Exception(f'PDF处理失败: {error_msg}')

            # 获取任务ID和结果
            if 'data' not in result or 'task_id' not in result['data']:
                raise Exception("MinerU返回结果缺少task_id")

            mineru_task_id = result['data']['task_id']
            logger.info(f"MinerU任务ID: {mineru_task_id}")

            if 'full_zip_url' not in result['data']:
                raise Exception("MinerU返回结果缺少full_zip_url")

            zip_url = result['data']['full_zip_url']
            logger.info(f"ZIP文件下载地址: {zip_url}")

            # 下载结果
            zip_filename = f"mineru_result_{mineru_task_id}.zip"
            zip_path = os.path.join(task_work_dir, zip_filename)

            # 下载或复制ZIP文件
            try:
                logger.info(f"开始处理ZIP文件: {zip_url}")

                if zip_url.startswith('file://'):
                    source_path = zip_url[7:]
                    logger.info(f"复制本地文件: {source_path} -> {zip_path}")
                    if not os.path.exists(source_path):
                        raise Exception(f"源文件不存在: {source_path}")
                    import shutil
                    shutil.copy2(source_path, zip_path)
                else:
                    logger.info(f"下载远程ZIP文件: {zip_url}")
                    # 禁用代理，直接连接
                    response = requests.get(zip_url, timeout=300, proxies={'http': None, 'https': None})
                    if response.status_code != 200:
                        raise Exception(f"下载ZIP文件失败，状态码: {response.status_code}")
                    with open(zip_path, 'wb') as f:
                        f.write(response.content)
                    logger.info(f"ZIP文件已保存到: {zip_path}")

            except Exception as e:
                logger.error(f"处理结果文件失败: {e}")
                raise

            # 解压ZIP文件
            try:
                logger.info(f"开始解压ZIP文件: {zip_path}")
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(task_work_dir)
                    logger.info(f"ZIP文件已解压到: {task_work_dir}")
            except Exception as e:
                logger.error(f"解压文件失败: {e}")
                raise

            # 查找markdown文件
            md_file = None
            for root, dirs, files in os.walk(task_work_dir):
                for file in files:
                    if file.endswith('.md') and mineru_task_id in file:
                        md_file = os.path.join(root, file)
                        logger.info(f"找到markdown文件: {md_file}")
                        break
                if md_file:
                    break

            if not md_file:
                # 查找任何md文件
                for root, dirs, files in os.walk(task_work_dir):
                    for file in files:
                        if file.endswith('.md'):
                            md_file = os.path.join(root, file)
                            logger.info(f"找到md文件: {md_file}")
                            break
                    if md_file:
                        break

            # 创建docx文件
            docx_filename = f"{os.path.splitext(unique_filename)[0]}.docx"
            docx_path = os.path.join(pdf_output_dir, docx_filename)

            if md_file:
                # 读取提取的文本内容
                with open(md_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                logger.info(f"成功读取内容文件，长度: {len(content)} 字符")

                # 使用翻译功能生成双语Word文档
                try:
                    from app.utils.document_generator import translate_markdown_to_bilingual_doc
                    
                    # 将语言代码映射
                    lang_mapping = {
                        'EN': 'en', 'en': 'en',
                        'ZH': 'zh', 'zh': 'zh',
                        'JA': 'ja', 'ja': 'ja'
                    }
                    source_language = lang_mapping.get(source_lang, 'en')
                    target_language = lang_mapping.get(target_lang, 'zh')

                    logger.info(f"开始翻译，源语言={source_language}, 目标语言={target_language}")
                    
                    # === PDF图片OCR识别和翻译 ===
                    ocr_results = []
                    if enable_image_ocr:
                        logger.info("=" * 60)
                        logger.info("开始PDF图片OCR识别和翻译")
                        logger.info("=" * 60)
                        try:
                            from app.function.image_ocr.ocr_controller import process_markdown_images_ocr_and_translate
                            
                            markdown_dir = os.path.dirname(md_file)
                            logger.info(f"Markdown文件目录: {markdown_dir}")
                            logger.info(f"源语言: {source_language}, 目标语言: {target_language}")
                            
                            # 调用OCR处理函数
                            ocr_results = process_markdown_images_ocr_and_translate(
                                markdown_content=content,
                                markdown_dir=markdown_dir,
                                target_language=target_language,
                                source_language=source_language
                            )
                            
                            if ocr_results:
                                logger.info(f"  OCR处理完成，共处理 {len(ocr_results)} 个图片")
                                for i, result in enumerate(ocr_results):
                                    if result.get('success'):
                                        logger.info(f"  图片 {i+1}: {os.path.basename(result['image_path'])}")
                                        logger.info(f"    OCR文本长度: {len(result.get('ocr_text_combined', ''))}")
                                        logger.info(f"    翻译文本长度: {len(result.get('translation_text_combined', ''))}")
                            else:
                                logger.info("未找到需要OCR处理的图片")
                                
                        except Exception as ocr_error:
                            logger.error(f"PDF图片OCR处理失败: {ocr_error}")
                            logger.exception("OCR错误详情")
                            # OCR失败不影响正常翻译流程
                    else:
                        logger.info("未启用图片OCR功能")
                    
                    ok = translate_markdown_to_bilingual_doc(
                        content,
                        docx_path,
                        source_language=source_language,
                        target_language=target_language,
                        image_base_dir=os.path.dirname(md_file),
                        custom_translations=custom_translations,
                        image_ocr_results=ocr_results  # 传递OCR结果
                    )
                    
                    if ok:
                        logger.info("翻译并生成Word文档成功")
                    else:
                        raise Exception("翻译生成Word文档失败")
                        
                except Exception as e:
                    logger.error(f"翻译过程中出错: {e}")
                    raise
            else:
                # 未找到md文件，创建提示文档
                from docx import Document
                doc = Document()
                doc.add_heading('PDF处理结果', 1)
                doc.add_paragraph('未能从PDF中提取到文本内容，请检查原始PDF文件是否包含可提取的文本。')
                doc.add_paragraph(f'原始文件名: {original_filename}')
                doc.add_paragraph(f'处理时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                doc.save(docx_path)
                logger.info("创建了包含提示信息的文档")

            # 检查文件是否存在
            if not os.path.exists(docx_path):
                raise Exception(f"翻译后的文件不存在: {docx_path}")

            # 获取文件大小
            file_size = os.path.getsize(docx_path)
            logger.info(f"文件大小: {file_size} 字节")

            # 保存到数据库
            from app import db
            from app.models import UploadRecord
            
            record = UploadRecord(
                filename=original_filename,
                stored_filename=docx_filename,
                file_path=pdf_output_dir,
                user_id=user_id,
                file_size=file_size,
                status='completed'
            )
            
            db.session.add(record)
            db.session.commit()
            logger.info(f"上传记录已保存到数据库，记录ID: {record.id}")

            # 更新任务状态到缓存
            base_name = os.path.splitext(original_filename)[0] if original_filename else os.path.splitext(docx_filename)[0]
            safe_base = secure_filename(base_name) if base_name else 'translated'
            download_name = f"translated_{safe_base}.docx"
            with pdf_task_lock:
                pdf_task_status_cache[task_id] = {
                    'status': 'completed',
                    'filename': docx_filename,
                    'stored_filename': docx_filename,
                    'original_filename': original_filename,
                    'download_name': download_name,
                    'message': '翻译完成'
                }

            logger.info(f"PDF翻译任务完成: {task_id}")
            return True

        except Exception as e:
            logger.error(f"PDF翻译任务失败: {e}")
            import traceback
            logger.error(f"错误详情: {traceback.format_exc()}")
            
            # 更新任务状态为失败
            with pdf_task_lock:
                pdf_task_status_cache[task_id] = {
                    'status': 'failed',
                    'error': str(e),
                    'message': '翻译失败'
                }
            
            raise


@main.route("/")
@login_required
def index():
    return render_template("main/index.html", user=current_user)


@main.route("/dashboard")
@login_required
def dashboard():
    return redirect(url_for("main.index"))


@main.route("/index")
@login_required
def index_page():
    return render_template("main/index.html", user=current_user)


@main.route("/page1")
@login_required
def page1():
    return render_template("main/page1.html", user=current_user)


@main.route("/page2")
@login_required
def page2():
    return render_template("main/page2.html", user=current_user)


# 允许的文件扩展名和大小限制
ALLOWED_EXTENSIONS = {"ppt", "pptx"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def get_unique_filename(filename):
    """生成唯一的文件名"""
    ext = filename.rsplit(".", 1)[1].lower()
    return f"{uuid.uuid4().hex}.{ext}"


def custom_filename(name):
    # 移除危险的路径字符，仅保留基本合法字符 + 中文
    name = re.sub(r'[\\/:"*?<>|]+', "_", name)  # 替换非法字符
    return name


@main.route('/upload', methods=['POST'])
@login_required
def upload_file():
    try:
        # 验证用户是否登录
        if not current_user.is_authenticated:
            return jsonify({"code": 403, "msg": "用户未登录"}), 403

        # 获取表单数据
        user_language = request.form.get("source_language", "English")
        target_language = request.form.get("target_language", "Chinese")
        bilingual_translation = request.form.get("bilingual_translation", "paragraph_up")
        select_page = request.form.getlist("select_page")
        model = request.form.get("model", "qwen")
        enable_text_splitting = request.form.get("enable_text_splitting", "False")  # 字符串: "False" 或 "True_spliting"
        enable_uno_conversion = request.form.get("enable_uno_conversion", "True").lower() == "true"

        # 获取选中的词汇表ID
        selected_vocabulary = request.form.get("selected_vocabulary", "")
        vocabulary_ids = []
        if selected_vocabulary:
            try:
                vocabulary_ids = [int(x.strip()) for x in selected_vocabulary.split(",") if x.strip()]
                logger.info(f"接收到词汇表ID: {vocabulary_ids}")
            except ValueError as e:
                logger.error(f"词汇表ID解析失败: {selected_vocabulary}, 错误: {str(e)}")
                vocabulary_ids = []

        # 记录接收到的参数
        logger.info(f"接收到的翻译参数:")
        logger.info(f"  - 源语言: {user_language}")
        logger.info(f"  - 目标语言: {target_language}")
        logger.info(f"  - 双语翻译: {bilingual_translation}")
        logger.info(f"  - 模型: {model}")
        logger.info(f"  - 文本分割: {enable_text_splitting}")
        logger.info(f"  - UNO转换: {enable_uno_conversion}")
        logger.info(f"  - 选择页面: {select_page}")
        logger.info(f"  - 词汇表数量: {len(vocabulary_ids)}")

        # 转换select_page为整数列表
        if select_page and select_page[0]:
            try:
                select_page = [int(x) for x in select_page[0].split(",")]
                logger.info(f"  用户选择的页面: {select_page}")
            except Exception as e:
                logger.error(f"  页面选择参数解析失败: {select_page}, 错误: {str(e)}")
                select_page = []
        else:
            logger.info(f"  没有选择页面，将翻译所有页面")
            select_page = []

        # 构建自定义翻译词典
        custom_translations = {}
        if vocabulary_ids:
            try:
                # 查询词汇表数据（包含权限检查）
                translations = Translation.query.filter(
                    Translation.id.in_(vocabulary_ids),
                    db.or_(
                        db.and_(Translation.user_id == current_user.id, Translation.is_public == False),
                        Translation.is_public == True,
                    ),
                ).all()

                logger.info(f"从数据库查询到 {len(translations)} 个词汇条目")

                # 根据翻译方向构建词典
                for trans in translations:
                    source_text = None
                    target_text = None

                    # 根据语言方向映射源文本和目标文本
                    if user_language == "English" and target_language == "Chinese":
                        source_text = trans.english
                        target_text = trans.chinese
                    elif user_language == "Chinese" and target_language == "English":
                        source_text = trans.chinese
                        target_text = trans.english
                    elif user_language == "English" and target_language == "Dutch":
                        source_text = trans.english
                        target_text = trans.dutch
                    elif user_language == "Dutch" and target_language == "English":
                        source_text = trans.dutch
                        target_text = trans.english
                    elif user_language == "Chinese" and target_language == "Dutch":
                        source_text = trans.chinese
                        target_text = trans.dutch
                    elif user_language == "Dutch" and target_language == "Chinese":
                        source_text = trans.dutch
                        target_text = trans.chinese

                    # 添加到词典（确保源文本和目标文本都存在且不为空）
                    if source_text and target_text and source_text.strip() and target_text.strip():
                        custom_translations[source_text.strip()] = target_text.strip()

                logger.info(f"构建自定义词典完成，包含 {len(custom_translations)} 个词汇对")
                logger.info(
                    f"词典示例: {dict(list(custom_translations.items())[:3])}..." if custom_translations else "词典为空")

            except Exception as e:
                logger.error(f"构建自定义词典失败: {str(e)}")
                custom_translations = {}

        # 其他参数处理
        stop_words_input = request.form.get("stop_words", "")
        stop_words = [word.strip() for word in stop_words_input.split("\n") if word.strip()]

        custom_translations_input = request.form.get("custom_translations", "")
        # 合并用户输入的翻译和词汇表翻译
        for line in custom_translations_input.split("\n"):
            line = line.strip()
            if not line:
                continue
            parts = line.split("->")
            if len(parts) == 2:
                eng, chi = parts[0].strip(), parts[1].strip()
                custom_translations[eng] = chi

        # 获取上传的文件
        file = request.files.get("file")

        if not file:
            return jsonify({"code": 400, "msg": "请选择文件上传"}), 400

        # 检查文件名和类型
        if not file.filename or not allowed_file(file.filename):
            return jsonify({"code": 400, "msg": "不支持的文件类型"}), 400

        # 检查文件大小
        file.seek(0, 2)  # 移动到文件末尾
        file_size = file.tell()  # 获取文件大小
        file.seek(0)  # 重置文件指针

        if file_size > MAX_FILE_SIZE:
            return jsonify({"code": 400, "msg": f"文件大小超过限制 ({MAX_FILE_SIZE/1024/1024}MB)"}), 400

        # 创建用户上传目录
        upload_folder = current_app.config["UPLOAD_FOLDER"]
        user_upload_dir = os.path.join(upload_folder, f"user_{current_user.id}")
        os.makedirs(user_upload_dir, exist_ok=True)

        # 生成安全的文件名
        original_filename = custom_filename(file.filename)

        # 创建语言名称到语言代码的映射
        language_map = {"English": "en", "Chinese": "zh", "Dutch": "nl"}

        # 获取源语言和目标语言的代码
        source_lang_code = language_map.get(user_language, user_language)
        target_lang_code = language_map.get(target_language, target_language)

        # 生成新的文件名格式：源语言_目标语言_源文件名.pptx
        name_without_ext, ext = os.path.splitext(original_filename)
        new_filename = f"{source_lang_code}_{target_lang_code}_{name_without_ext}{ext}"

        stored_filename = get_unique_filename(new_filename)
        file_path = os.path.join(user_upload_dir, stored_filename)

        try:
            # 保存PPT文件
            file.save(file_path)

            # 创建上传记录，使用新的文件名
            record = UploadRecord(
                user_id=current_user.id,
                filename=new_filename,  # 使用新的文件名格式
                stored_filename=stored_filename,
                file_path=user_upload_dir,
                file_size=file_size,
                status="pending",
            )

            db.session.add(record)
            db.session.commit()

            # 添加翻译任务到队列
            priority = 0  # 默认优先级

            # 记录传递给任务队列的参数
            logger.info(f"传递给任务队列的参数:")
            logger.info(f"  - 文件路径: {file_path}")
            logger.info(f"  - 模型: {model}")
            logger.info(f"  - 文本分割: {enable_text_splitting}")
            logger.info(f"  - UNO转换: {enable_uno_conversion}")
            logger.info(f"  - 自定义词典条目数: {len(custom_translations)}")

            queue_position = translation_queue.add_task(
                user_id=current_user.id,
                user_name=current_user.username,
                file_path=file_path,
                select_page=select_page,
                source_language=user_language,
                target_language=target_language,
                bilingual_translation=bilingual_translation,
                priority=priority,
                model=model,
                enable_text_splitting=enable_text_splitting,
                enable_uno_conversion=enable_uno_conversion,
                custom_translations=custom_translations,  # 传递自定义词典
            )

            return jsonify(
                {
                    "code": 200,
                    "msg": "文件上传成功，已加入翻译队列",
                    "queue_position": queue_position,
                    "record_id": record.id,
                }
            )

        except Exception as e:
            # 清理已上传的文件
            if os.path.exists(file_path):
                os.remove(file_path)

            # 回滚数据库事务
            db.session.rollback()

            logger.error(f"文件上传失败: {str(e)}")
            return jsonify({"code": 500, "msg": f"文件上传失败: {str(e)}"}), 500

    except Exception as e:
        logger.error(f"处理上传请求失败: {str(e)}")
        return jsonify({"code": 500, "msg": f"处理上传请求失败: {str(e)}"}), 500


def process_queue(app, stop_words_list, custom_translations, source_language, target_language, bilingual_translation):
    """
    处理翻译队列的函数

    注意：此函数已被 EnhancedTranslationQueue 类的 _processor_loop 方法取代，
    不再被主动调用。保留此函数仅用于兼容旧代码。
    新的任务处理逻辑在 app/utils/enhanced_task_queue.py 中实现。
    """
    while True:
        task = translation_queue.start_next_task()
        if not task:
            time.sleep(1)  # 如果没有任务，等待1秒
            continue

        # 创建应用上下文
        with app.app_context():
            # try:
            # 执行翻译
            process_presentation(
                task["file_path"],
                stop_words_list,
                custom_translations,
                task["select_page"],
                source_language,
                target_language,
                bilingual_translation,
                model=task.get("model", "qwen"),
                enable_text_splitting=task.get("enable_text_splitting", "False"),
            )

            set_textbox_autofit(task["file_path"])

            translation_queue.complete_current_task(success=True)

            # 更新数据库记录状态
            record = UploadRecord.query.filter_by(
                user_id=task["user_id"],
                file_path=os.path.dirname(task["file_path"]),
                stored_filename=os.path.basename(task["file_path"]),
            ).first()

            if record:
                record.status = "completed"
                db.session.commit()

            # except Exception as e:
            #     print(f"Translation error: {str(e)}")
            #     translation_queue.complete_current_task(success=False, error=str(e))

            # 更新数据库记录状态
            if "record" in locals() and record:
                record.status = "failed"
                try:
                    db.session.commit()
                except:
                    db.session.rollback()
        # finally:
        #     # 确保会话被正确清理
        #     db.session.remove()


@main.route("/task_status")
@login_required
def get_task_status():
    """获取当前用户的任务状态"""
    status = translation_queue.get_task_status_by_user(current_user.id)
    if status:
        # 转换日志格式以便前端显示
        if "recent_logs" in status:
            formatted_logs = []
            for log in status["recent_logs"]:
                formatted_logs.append(
                    {
                        "timestamp": datetime_to_isoformat(log["timestamp"]) if log["timestamp"] else "",
                        "message": log["message"],
                        "level": log["level"],
                    }
                )
            status["recent_logs"] = formatted_logs

        # 使用ISO格式化时间戳
        for key in ["created_at", "started_at", "completed_at"]:
            if key in status and status[key]:
                status[key] = datetime_to_isoformat(status[key])

        return jsonify(status)
    return jsonify({"status": "no_task"})


@main.route('/api/pdf_task_status')
@login_required
def get_pdf_task_status():
    """获取当前用户的PDF翻译任务状态"""
    try:
        # 从session中获取任务ID
        task_id = session.get('pdf_task_id')
        
        if not task_id:
            return jsonify({'status': 'no_task'})
        
        # 从缓存中获取任务状态
        with pdf_task_lock:
            task_data = pdf_task_status_cache.get(task_id)
        
        if not task_data:
            # 如果缓存中没有，检查session中是否有初始状态
            task_status = session.get('pdf_task_status', 'no_task')
            if task_status == 'waiting':
                original_name = session.get('pdf_original_filename', '')
                response = {
                    'status': 'waiting',
                    'message': '任务正在排队...'
                }
                if original_name:
                    response['original_filename'] = original_name
                return jsonify(response)
            return jsonify({'status': 'no_task'})
        
        response = {'status': task_data['status']}
        if 'message' in task_data and task_data['message']:
            response['message'] = task_data['message']
        
        if task_data['status'] == 'completed':
            stored_filename = task_data.get('stored_filename', '')
            original_filename = task_data.get('original_filename') or session.get('pdf_original_filename', '')
            download_name = task_data.get('download_name')
            if not download_name:
                base_name = os.path.splitext(original_filename or stored_filename)[0]
                safe_base = secure_filename(base_name) if base_name else 'translated'
                download_name = f"translated_{safe_base}.docx"

            response['filename'] = task_data.get('filename', stored_filename)
            response['stored_filename'] = stored_filename
            response['original_filename'] = original_filename
            response['download_name'] = download_name

            # 清理session中的任务状态
            session.pop('pdf_task_id', None)
            session.pop('pdf_task_status', None)
            session.pop('pdf_original_filename', None)
            # 延迟清理缓存（5秒后），确保前端能接收到完成状态
            def cleanup_cache():
                time.sleep(5)
                with pdf_task_lock:
                    pdf_task_status_cache.pop(task_id, None)
            import threading
            threading.Thread(target=cleanup_cache, daemon=True).start()
        elif task_data['status'] == 'failed':
            response['error'] = task_data.get('error', '翻译失败')
            # 清理session中的任务状态
            session.pop('pdf_task_id', None)
            session.pop('pdf_task_status', None)
            session.pop('pdf_original_filename', None)
            # 延迟清理缓存（5秒后）
            def cleanup_cache():
                time.sleep(5)
                with pdf_task_lock:
                    pdf_task_status_cache.pop(task_id, None)
            import threading
            threading.Thread(target=cleanup_cache, daemon=True).start()
        elif task_data['status'] == 'processing':
            response['message'] = task_data.get('message', '正在翻译中...')
        elif task_data['status'] == 'waiting':
            response['message'] = task_data.get('message', '任务正在排队...')
            
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"获取PDF任务状态失败: {e}")
        return jsonify({'status': 'no_task'})


@main.route('/queue_status')
@login_required
def get_queue_status():
    """获取翻译队列状态信息"""
    try:
        # 获取队列统计信息
        queue_stats = translation_queue.get_queue_stats()

        # 添加详细的任务信息
        active_tasks = queue_stats.get("processing", 0)  # 修正键名
        waiting_tasks = queue_stats.get("waiting", 0)
        max_concurrent = queue_stats.get("max_concurrent", 10)

        detailed_stats = {
            "max_concurrent_tasks": max_concurrent,
            "active_tasks": active_tasks,
            "waiting_tasks": waiting_tasks,
            "total_tasks": queue_stats.get("total", 0),
            "completed_tasks": queue_stats.get("completed", 0),
            "failed_tasks": queue_stats.get("failed", 0),
            "available_slots": max(0, max_concurrent - active_tasks),
            "queue_full": (active_tasks + waiting_tasks) >= max_concurrent,
            "system_status": "normal" if (active_tasks + waiting_tasks) < max_concurrent else "busy",
        }

        # 如果是管理员，提供更多详细信息
        if current_user.is_administrator():
            detailed_stats["admin_info"] = {
                "processor_running": translation_queue.running,
                "task_timeout": translation_queue.task_timeout,
                "retry_times": translation_queue.retry_times,
            }

        return jsonify(detailed_stats)

    except Exception as e:
        logger.error(f"获取队列状态失败: {str(e)}")
        return (
            jsonify(
                {
                    "error": "获取队列状态失败",
                    "max_concurrent_tasks": 10,
                    "active_tasks": 0,
                    "waiting_tasks": 0,
                    "total_tasks": 0,
                    "available_slots": 10,
                    "queue_full": False,
                    "system_status": "unknown",
                }
            ),
            500,
        )


@main.route("/history")
@login_required
def get_history():
    try:
        # 只返回状态为 completed 的记录
        records = (
            UploadRecord.query.filter_by(user_id=current_user.id, status="completed")
            .order_by(UploadRecord.upload_time.desc())
            .all()
        )

        history_records = []
        for record in records:
            # 检查文件是否仍然存在
            file_exists = os.path.exists(os.path.join(record.file_path, record.stored_filename))

            # 使用ISO格式返回时间，让前端正确处理时区
            upload_time = datetime_to_isoformat(record.upload_time)

            # 直接使用数据库中存储的文件名
            history_records.append(
                {
                    "id": record.id,
                    "filename": record.filename,  # 使用数据库中存储的文件名
                    "file_size": record.file_size,
                    "upload_time": upload_time,
                    "status": record.status,
                    "file_exists": file_exists,
                }
            )

        return jsonify(history_records)

    except Exception as e:
        print(f"History error: {str(e)}")
        return jsonify({"status": "error", "message": "获取历史记录失败"}), 500


@main.route("/download/<int:record_id>")
@login_required
def download_file(record_id):
    try:
        # 获取上传记录
        record = UploadRecord.query.get_or_404(record_id)

        # 验证用户权限
        if record.user_id != current_user.id:
            return jsonify({"error": "无权访问此文件"}), 403

        # 检查文件是否存在
        file_path = os.path.join(record.file_path, record.stored_filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "文件不存在"}), 404

        # 添加调试信息
        print(f"Downloading file: {file_path}")
        print(f"Original filename: {record.filename}")
        file_path = os.path.abspath(file_path)
        return send_file(file_path, as_attachment=True, download_name=record.filename)
    except Exception as e:
        print(f"Download error: {str(e)}")
        return jsonify({"error": f"下载失败: {str(e)}"}), 500


@main.route("/delete/<int:record_id>", methods=["DELETE"])
@login_required
def delete_file(record_id):
    try:
        # 获取上传记录
        record = UploadRecord.query.get_or_404(record_id)

        # 验证用户权限
        if record.user_id != current_user.id:
            return jsonify({"error": "无权删除此文件"}), 403

        try:
            # 删除物理文件
            file_path = os.path.join(record.file_path, record.stored_filename)
            if os.path.exists(file_path):
                os.remove(file_path)

            # 删除数据库记录
            db.session.delete(record)
            db.session.commit()

            return jsonify({"message": "文件删除成功"})

        except Exception as e:
            db.session.rollback()
            print(f"Delete error: {str(e)}")
            return jsonify({"error": f"删除失败: {str(e)}"}), 500

    except Exception as e:
        print(f"Delete error: {str(e)}")
        return jsonify({"error": f"删除失败: {str(e)}"}), 500


@main.route("/translate")
@login_required
def translate():
    return render_template("main/translate.html", user=current_user)


@main.route('/pdf_translate')
@login_required
def pdf_translate():
    """PDF翻译页面"""
    return render_template("main/pdf_translate.html")


@main.route("/batch_process")
@login_required
def batch_process():
    return render_template("main/batch_process.html", user=current_user)


@main.route("/settings")
@login_required
def settings():
    return render_template("main/settings.html", user=current_user)


@main.route('/pdf_translation')
@login_required
def pdf_translation():
    return render_template("main/pdf_translation.html", user=current_user)


@main.route('/dictionary')
@login_required
def dictionary():
    return render_template("main/dictionary.html", user=current_user)


@main.route("/file_search")
@login_required
def file_search():
    return render_template("main/file_search.html", user=current_user)


@main.route("/account_settings")
@login_required
def account_settings():
    return render_template("main/account_settings.html", user=current_user)


@main.route("/registration_approval")
@login_required
def registration_approval():
    if not current_user.is_administrator():
        flash("没有权限访问此页面")
        return redirect(url_for("main.index"))
    return render_template("main/registration_approval.html")


# @main.route('/sso_management')
# @login_required
# def sso_management():
#     """SSO管理页面"""
#     if not current_user.is_administrator():
#         flash('没有权限访问此页面')
#         return redirect(url_for('main.index'))
#     return render_template('main/sso_management.html')


@main.route("/api/registrations")
@login_required
def get_registrations():
    if not current_user.is_administrator():
        return jsonify({"error": "没有权限访问"}), 403

    page = request.args.get("page", 1, type=int)
    status = request.args.get("status", "all")
    per_page = 10

    query = User.query
    if status != "all":
        query = query.filter_by(status=status)

    pagination = query.order_by(User.register_time.desc()).paginate(page=page, per_page=per_page, error_out=False)

    return jsonify(
        {
            "registrations": [
                {
                    "id": user.id,
                    "username": user.username,
                    "status": user.status,
                    "register_time": datetime_to_isoformat(user.register_time) if user.register_time else None,
                    "approve_user": user.approve_user.username if user.approve_user else None,
                    "approve_time": datetime_to_isoformat(user.approve_time) if user.approve_time else None,
                }
                for user in pagination.items
            ],
            "total_pages": pagination.pages,
            "current_page": page,
            "total": pagination.total,
        }
    )


@main.route("/api/users")
@login_required
def get_users():
    if not current_user.is_administrator():
        return jsonify({"error": "没有权限访问"}), 403

    page = request.args.get("page", 1, type=int)
    status = request.args.get("status", "all")
    per_page = 10

    query = User.query.filter(User.status.in_(["approved", "disabled"]))
    if status != "all":
        query = query.filter_by(status=status)

    pagination = query.order_by(User.register_time.desc()).paginate(page=page, per_page=per_page, error_out=False)

    return jsonify(
        {
            "users": [
                {
                    "id": user.id,
                    "username": user.username,
                    "status": user.status,
                    "register_time": datetime_to_isoformat(user.register_time) if user.register_time else None,
                }
                for user in pagination.items
            ],
            "total_pages": pagination.pages,
            "current_page": page,
            "total": pagination.total,
        }
    )


@main.route("/api/registrations/<int:id>/approve", methods=["POST"])
@login_required
def approve_registration(id):
    if not current_user.is_administrator():
        return jsonify({"error": "没有权限进行此操作"}), 403

    user = User.query.get_or_404(id)
    if user.status != "pending":
        return jsonify({"error": "该用户已被审批"}), 400

    try:
        user.status = "approved"
        user.approve_time = datetime.now(pytz.timezone("Asia/Shanghai"))
        user.approve_user_id = current_user.id
        db.session.commit()
        return jsonify({"message": "审批成功"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@main.route("/api/registrations/<int:id>/reject", methods=["POST"])
@login_required
def reject_registration(id):
    if not current_user.is_administrator():
        return jsonify({"error": "没有权限进行此操作"}), 403

    user = User.query.get_or_404(id)
    if user.status != "pending":
        return jsonify({"error": "该用户已被审批"}), 400

    try:
        user.status = "rejected"
        user.approve_time = datetime.now(pytz.timezone("Asia/Shanghai"))
        user.approve_user_id = current_user.id
        db.session.commit()
        return jsonify({"message": "已拒绝申请"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@main.route("/api/users/<int:id>/disable", methods=["POST"])
@login_required
def disable_user(id):
    if not current_user.is_administrator():
        return jsonify({"error": "没有权限进行此操作"}), 403

    user = User.query.get_or_404(id)
    if user.status != "approved":
        return jsonify({"error": "该用户无法被禁用"}), 400

    try:
        user.status = "disabled"
        db.session.commit()
        return jsonify({"message": "用户已禁用"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@main.route("/api/users/<int:id>/enable", methods=["POST"])
@login_required
def enable_user(id):
    if not current_user.is_administrator():
        return jsonify({"error": "没有权限进行此操作"}), 403

    user = User.query.get_or_404(id)
    if user.status != "disabled":
        return jsonify({"error": "该用户无法被启用"}), 400

    try:
        user.status = "approved"
        db.session.commit()
        return jsonify({"message": "用户已启用"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


# 词库管理API路由
@main.route("/api/translations", methods=["GET"])
@login_required
def get_translations():
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 10, type=int)  # 添加per_page参数支持
    search = request.args.get("search", "")
    # Add filter for public/private translations
    visibility = request.args.get("visibility", "private")  # private, public, all

    if visibility == "private":
        # 只查询当前用户的私有翻译数据
        query = Translation.query.filter(Translation.user_id == current_user.id, Translation.is_public == False)
    elif visibility == "public":
        # 只查询公共的翻译数据
        query = Translation.query.filter_by(is_public=True)
    else:  # all 或其他值，默认为all
        # 查询当前用户的所有私有数据和所有公共数据
        query = Translation.query.filter(
            db.or_(
                db.and_(Translation.user_id == current_user.id, Translation.is_public == False),
                Translation.is_public == True,
            )
        )

    if search:
        query = query.filter(
            db.or_(
                Translation.english.ilike(f"%{search}%"),
                Translation.chinese.ilike(f"%{search}%"),
                Translation.dutch.ilike(f"%{search}%"),
                Translation.category.ilike(f"%{search}%"),
            )
        )

    pagination = query.order_by(Translation.id.desc()).paginate(page=page, per_page=per_page, error_out=False)

    translations_data = []
    for item in pagination.items:
        translation_dict = {
            "id": item.id,
            "english": item.english,
            "chinese": item.chinese,
            "dutch": item.dutch,
            "category": item.category,
            "created_at": datetime_to_isoformat(item.created_at),
            "is_public": item.is_public,
            "user_id": item.user_id,
        }
        # Add user info for display
        if item.user:
            translation_dict["user"] = {"id": item.user.id, "username": item.user.username}
        translations_data.append(translation_dict)

    return jsonify(
        {
            "translations": translations_data,
            "total_pages": pagination.pages,
            "current_page": page,
            "total_items": pagination.total,
        }
    )


@main.route("/api/translations/categories", methods=["GET"])
@login_required
def get_translation_categories():
    """获取所有已存在的分类列表（去重，按字母排序）"""
    try:
        from app.models.translation import Translation as TranslationModel

        categories_set = set()
        # 仅提取有值的分类
        for row in db.session.query(TranslationModel.category).filter(TranslationModel.category.isnot(None)).all():
            value = row[0]
            if not value:
                continue
            # 支持分号分隔的多分类
            for part in value.split(";"):
                name = part.strip()
                if name:
                    categories_set.add(name)
        categories = sorted(categories_set, key=lambda x: x.lower())
        return jsonify({"categories": categories})
    except Exception as e:
        logger.error(f"获取分类失败: {e}")
        return jsonify({"categories": []}), 200


@main.route("/api/translations", methods=["POST"])
@login_required
def add_translation():
    data = request.get_json()
    english = data.get("english")
    chinese = data.get("chinese")
    dutch = data.get("dutch")
    category = data.get("category")  # Single category field
    is_public = data.get("is_public", False)

    if not english or not chinese:
        return jsonify({"error": "英文和中文翻译都是必填的"}), 400

    # Build query based on whether it's a public or private translation
    if is_public and current_user.is_administrator():
        # For public translations, check against all public translations
        existing = Translation.query.filter_by(english=english, is_public=True).first()
    else:
        # For private translations, check only against current user's translations
        is_public = False  # Ensure non-admin users can't add public translations
        existing = Translation.query.filter_by(user_id=current_user.id, english=english).first()

    if existing:
        return jsonify({"error": "该英文翻译已存在于词库中"}), 400

    try:
        translation = Translation(
            english=english,
            chinese=chinese,
            dutch=dutch,
            category=category,
            is_public=is_public,
            user_id=current_user.id,  # Always set user_id, even for public translations
        )
        db.session.add(translation)
        db.session.commit()

        return jsonify(
            {
                "message": "添加成功",
                "translation": {
                    "id": translation.id,
                    "english": translation.english,
                    "chinese": translation.chinese,
                    "dutch": translation.dutch,
                    "category": translation.category,
                    "is_public": translation.is_public,
                    "created_at": datetime_to_isoformat(translation.created_at),
                },
            }
        )
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@main.route("/api/translations/<int:id>", methods=["DELETE"])
@login_required
def delete_translation(id):
    translation = Translation.query.get_or_404(id)

    # 验证所有权 - users can only delete their own private translations
    # admins can delete public translations
    if translation.is_public:
        if not current_user.is_administrator():
            return jsonify({"error": "无权删除公共词库"}), 403
    else:
        if translation.user_id != current_user.id:
            return jsonify({"error": "无权删除此翻译"}), 403

    try:
        db.session.delete(translation)
        db.session.commit()
        return jsonify({"message": "删除成功"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@main.route("/api/translations/<int:id>", methods=["PUT"])
@login_required
def update_translation(id):
    translation = Translation.query.get_or_404(id)

    # 验证所有权 - users can only edit their own private translations
    # admins can edit public translations
    if translation.is_public:
        if not current_user.is_administrator():
            return jsonify({"error": "无权修改公共词库"}), 403
    else:
        if translation.user_id != current_user.id:
            return jsonify({"error": "无权修改此翻译"}), 403

    data = request.get_json()
    english = data.get("english")
    chinese = data.get("chinese")
    is_public = data.get("is_public", translation.is_public)  # Keep existing value if not provided

    # Only admins can change the public status
    if "is_public" in data and data["is_public"] != translation.is_public:
        if not current_user.is_administrator():
            return jsonify({"error": "无权修改词条的公共状态"}), 403

    if not english or not chinese:
        return jsonify({"error": "英文和中文翻译都是必填的"}), 400

    # 检查是否与其他翻译重复
    if translation.is_public or is_public:
        # For public translations, check against all public translations
        existing = Translation.query.filter(
            Translation.is_public == True, Translation.english == english, Translation.id != id
        ).first()
    else:
        # For private translations, check only against current user's translations
        existing = Translation.query.filter(
            Translation.user_id == current_user.id, Translation.english == english, Translation.id != id
        ).first()

    if existing:
        return jsonify({"error": "该英文翻译已存在于词库中"}), 400

    try:
        translation.english = english
        translation.chinese = chinese
        translation.dutch = data.get("dutch")
        translation.category = data.get("category")

        # Only admins can change public status
        if current_user.is_administrator() and "is_public" in data:
            translation.is_public = is_public

        db.session.commit()

        return jsonify(
            {
                "message": "更新成功",
                "translation": {
                    "id": translation.id,
                    "english": translation.english,
                    "chinese": translation.chinese,
                    "dutch": translation.dutch,
                    "category": translation.category,
                    "is_public": translation.is_public,
                    "created_at": datetime_to_isoformat(translation.created_at),
                },
            }
        )
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@main.route("/api/translations/stats", methods=["GET"])
@login_required
def get_translation_stats():
    """获取当前用户的词库统计信息"""
    try:
        total_count = Translation.query.filter_by(user_id=current_user.id).count()
        return jsonify({"total_translations": total_count})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@main.route("/api/train", methods=["POST"])
@login_required
def train_model():
    """使用当前用户的词库数据进行训练"""
    try:

        # Tokenizer()
        # # TODO: 实现模型训练逻辑，只使用当前用户的数据
        # train_model()
        translations = Translation.query.all()
        return jsonify({"message": "训练完成", "data_count": len(translations)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@main.route("/ingredient")
@login_required
def ingredient():
    return render_template("main/ingredient.html")


@main.route("/ingredient/upload")
@login_required
def ingredient_upload_page():
    if not current_user.is_administrator():
        abort(403)
    return render_template("main/ingredient_upload.html")


# 加载JSON数据
def load_data(json_path):
    with open(json_path, "r", encoding="UTF-8") as file:
        return json.load(file)


def extract_ingredient(s, ingredient):
    """提取匹配的成分"""
    ingredients = re.sub(r"(\(|\（)", ",", s)
    ingredients = re.sub(r"(\)|\）)", "", ingredients)
    ingredients = re.split(r"[、,，]", ingredients)
    ingredients = [ing.replace(" ", "") for ing in ingredients]
    # 去掉类似于"又名"、"以"、"记"等词
    cleaned_ingredient_list = [re.sub(r"(又名|以|记)", "", ing) for ing in ingredients]

    for i in cleaned_ingredient_list:
        if ingredient in i:
            return i
    return None


def clean_food_name(food_name):
    """清理食品名称"""
    return re.sub(r"备案入.*", "", food_name)


@main.route("/search", methods=["POST"])
@login_required
def search_ingredient():
    # print(request.form['query'])
    # 临时返回空结果，直到实现完整的搜索功能
    return jsonify([])


@main.route("/ingredient/download", methods=["POST"])
@login_required
def download_ingredient_file():
    # print(request.form['file_path'])
    # 临时返回错误，直到实现完整的下载功能
    return jsonify({"error": "功能暂未实现"}), 500


# 允许的PDF文件扩展名
PDF_ALLOWED_EXTENSIONS = {"pdf"}


def allowed_pdf_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in PDF_ALLOWED_EXTENSIONS


@main.route("/pdf/<filename>")
@login_required
def get_pdf(filename):
    try:
        # 获取上传文件夹路径
        upload_folder = current_app.config["UPLOAD_FOLDER"]
        logger.info(f"PDF请求: {filename}, 上传文件夹: {upload_folder}")

        if not os.path.exists(upload_folder):
            logger.error(f"上传文件夹不存在: {upload_folder}")
            return jsonify({"error": "上传文件夹不存在"}), 404

        # 构建用户PDF目录路径
        user_pdf_dir = os.path.join(upload_folder, f"{current_user.username}_pdfs")
        logger.info(f"尝试从目录提供PDF: {user_pdf_dir}")

        if not os.path.exists(user_pdf_dir):
            # 尝试创建目录
            try:
                os.makedirs(user_pdf_dir, exist_ok=True)
                logger.info(f"创建了PDF目录: {user_pdf_dir}")
            except Exception as e:
                logger.error(f"无法创建PDF目录: {user_pdf_dir}, 错误: {str(e)}")
                return jsonify({"error": f"无法创建PDF目录: {str(e)}"}), 500

        # 构建完整的文件路径
        file_path = os.path.join(user_pdf_dir, filename)
        file_path = os.path.abspath(file_path)  # 转换为绝对路径
        logger.info(f"完整的PDF文件路径: {file_path}")

        if not os.path.exists(file_path):
            logger.error(f"PDF文件不存在: {file_path}")

            # 检查是否存在于其他可能的位置
            alt_paths = [
                os.path.join(upload_folder, filename),  # 直接在上传文件夹中
                os.path.join(upload_folder, "pdf", filename),  # 在pdf子文件夹中
                os.path.join(current_app.root_path, "static", "uploads", filename),  # 在静态文件夹中
            ]

            for alt_path in alt_paths:
                if os.path.exists(alt_path):
                    logger.info(f"在替代位置找到PDF文件: {alt_path}")
                    file_path = alt_path
                    break
            else:
                return jsonify({"error": "文件不存在"}), 404

        # 检查文件权限
        try:
            # 尝试打开文件进行读取测试
            with open(file_path, "rb") as f:
                f.read(1)  # 只读取1字节进行测试
            logger.info(f"文件权限检查通过: {file_path}")
        except PermissionError:
            logger.error(f"无法读取PDF文件(权限错误): {file_path}")
            # 尝试修改文件权限
            try:
                import stat

                os.chmod(file_path, stat.S_IRUSR | stat.S_IWUSR | stat.S_IRGRP | stat.S_IROTH)
                logger.info(f"已修改文件权限: {file_path}")
            except Exception as e:
                logger.error(f"无法修改文件权限: {str(e)}")
                return jsonify({"error": f"文件无法访问(权限错误): {str(e)}"}), 403
        except Exception as e:
            logger.error(f"文件读取测试失败: {str(e)}")
            return jsonify({"error": f"文件无法访问: {str(e)}"}), 403

        logger.info(f"准备提供PDF文件: {file_path}")
        try:
            # 使用安全的方式提供文件
            response = send_file(file_path, mimetype="application/pdf", as_attachment=False, download_name=filename)
            # 添加必要的安全头部
            response.headers["Access-Control-Allow-Origin"] = "*"
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"

            # 添加内容安全策略头部
            response.headers["Content-Security-Policy"] = (
                "default-src 'self'; img-src 'self' data:; style-src 'self' 'unsafe-inline'; script-src 'self' 'unsafe-inline'; object-src 'self'"
            )

            # 添加X-Content-Type-Options头部，防止MIME类型嗅探
            response.headers["X-Content-Type-Options"] = "nosniff"

            # 强制使用HTTPS
            if request.is_secure:
                response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"

            logger.info(f"PDF文件已成功提供: {file_path}")
            return response

        except Exception as e:
            logger.error(f"提供PDF文件时出错: {str(e)}")
            raise

    except Exception as e:
        logger.error(f"PDF提供错误: {str(e)}")
        return jsonify({"error": f"获取文件失败: {str(e)}"}), 500


@main.route("/save_annotations", methods=["POST"])
@login_required
def save_annotations():
    try:
        data = request.get_json()
        annotations = data.get("annotations", [])

        # 创建注释存储目录
        annotations_dir = os.path.join(current_app.config["UPLOAD_FOLDER"], f"{current_user.username}_annotations")

        if not os.path.exists(annotations_dir):
            os.makedirs(annotations_dir)

        # 保存注释到JSON文件
        filename = f"annotations_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        file_path = os.path.join(annotations_dir, filename)

        # 添加时间戳和用户信息
        annotation_data = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "user": current_user.username,
            "annotations": annotations,
        }

        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(annotation_data, f, ensure_ascii=False, indent=2)

        return jsonify({"message": "注释保存成功"})

    except Exception as e:
        print(f"Save annotations error: {str(e)}")
        return jsonify({"error": f"保存失败: {str(e)}"}), 500


@main.route("/get_annotations/<filename>")
@login_required
def get_annotations(filename):
    try:
        annotations_dir = os.path.join(current_app.config["UPLOAD_FOLDER"], f"{current_user.username}_annotations")

        file_path = os.path.join(annotations_dir, filename)

        if not os.path.exists(file_path):
            return jsonify({"error": "注释文件不存在"}), 404

        with open(file_path, "r", encoding="utf-8") as f:
            annotations = json.load(f)

        return jsonify(annotations)

    except Exception as e:
        print(f"Get annotations error: {str(e)}")
        return jsonify({"error": f"获取注释失败: {str(e)}"}), 500


@main.route("/get_annotation_files")
@login_required
def get_annotation_files():
    try:
        # 获取用户注释文件目录
        annotations_dir = os.path.join(current_app.config["UPLOAD_FOLDER"], f"{current_user.username}_annotations")

        if not os.path.exists(annotations_dir):
            return jsonify([])

        # 获取目录中的所有JSON文件
        files = []
        for filename in os.listdir(annotations_dir):
            if filename.endswith(".json"):
                file_path = os.path.join(annotations_dir, filename)
                files.append(
                    {
                        "filename": filename,
                        "created_time": datetime.fromtimestamp(os.path.getctime(file_path)).strftime(
                            "%Y-%m-%d %H:%M:%S"
                        ),
                    }
                )

        # 按创建时间降序排序
        files.sort(key=lambda x: x["created_time"], reverse=True)
        return jsonify(files)

    except Exception as e:
        print(f"Error getting annotation files: {str(e)}")
        return jsonify({"error": str(e)}), 500


@main.route("/api/users/sso")
@login_required
def get_sso_users():
    """获取SSO用户列表"""
    if not current_user.is_administrator():
        return jsonify({"error": "权限不足"}), 403

    try:
        # 查询所有SSO用户
        sso_users = User.query.filter(User.sso_provider.isnot(None)).all()

        users_data = []
        for user in sso_users:
            # 格式化时间
            last_login = format_datetime(user.last_login)
            register_time = format_datetime(user.register_time)

            users_data.append(
                {
                    "id": user.id,
                    "username": user.username,
                    "email": user.email or "",
                    "display_name": user.get_display_name(),
                    "first_name": user.first_name or "",
                    "last_name": user.last_name or "",
                    "sso_provider": user.sso_provider,
                    "sso_subject": user.sso_subject or "",
                    "status": user.status,
                    "role": user.role.name if user.role else "unknown",
                    "last_login": last_login,
                    "register_time": register_time,
                }
            )

        return jsonify(users_data)

    except Exception as e:
        logger.error(f"获取SSO用户列表失败: {e}")
        return jsonify({"error": f"获取用户列表失败: {str(e)}"}), 500


@main.route("/get_queue_status")
def get_detailed_queue_status():
    """获取详细的翻译队列状态（旧版API）"""
    username = session.get("username", "")
    if not username:
        return jsonify({"code": 403, "msg": "用户未登录"}), 403

    try:
        # 获取队列状态和统计信息
        status_info = translation_queue.get_queue_status()
        user_tasks = translation_queue.get_user_tasks(username)

        # 轮询用户任务以获取当前状态
        user_task_details = []
        for task in user_tasks:
            task_detail = {
                "task_id": task.task_id,
                "file_name": os.path.basename(task.file_path),
                "status": task.status,
                "progress": task.progress,
                "result": task.result,
                "error": task.error,
                "created_at": task.created_at.isoformat(),
                "started_at": task.started_at.isoformat() if task.started_at else None,
                "completed_at": task.completed_at.isoformat() if task.completed_at else None,
            }
            user_task_details.append(task_detail)

        return jsonify({"code": 200, "queue_status": status_info, "user_tasks": user_task_details})
    except Exception as e:
        logger.error(f"获取队列状态失败: {str(e)}")
        return jsonify({"code": 500, "msg": f"获取队列状态失败: {str(e)}"}), 500


@main.route("/cancel_task/<task_id>")
def cancel_task(task_id):
    """取消翻译任务"""
    username = session.get("username", "")
    if not username:
        return jsonify({"code": 403, "msg": "用户未登录"}), 403

    try:
        # 尝试取消任务
        result = translation_queue.cancel_task(task_id, username)
        if result:
            return jsonify({"code": 200, "msg": "任务已取消"})
        else:
            return jsonify({"code": 400, "msg": "取消任务失败，任务可能不存在或已经开始处理"}), 400
    except Exception as e:
        logger.error(f"取消任务失败: {str(e)}")
        return jsonify({"code": 500, "msg": f"取消任务失败: {str(e)}"}), 500


@main.route("/logs")
@login_required
def logs():
    """日志管理页面"""
    # 检查管理员权限
    if not current_user.is_administrator():
        flash("没有权限访问此页面", "error")
        return redirect(url_for("main.index"))
    return render_template("main/logs.html")


@main.route("/switch_language", methods=["POST"])
def switch_language():
    """处理语言切换请求"""
    try:
        data = request.get_json()
        language = data.get("language", "zh")

        # 验证语言代码
        if language not in ["zh", "en"]:
            return jsonify({"success": False, "message": "Invalid language code"}), 400

        # 在session中保存语言设置
        session["language"] = language

        return jsonify({"success": True, "message": "Language switched successfully", "language": language})

    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500



# ==================== 公开API端点（不需要认证） ====================
# 用于简单前端（html文件夹）的API端点


@main.route("/start_translation", methods=["POST"])
def start_translation():
    """启动PPT翻译任务（公开API，不需要认证）"""
    try:
        # 检查是否有文件
        if "file" not in request.files:
            return jsonify({"error": "没有文件"}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "没有选择文件"}), 400

        if not allowed_file(file.filename):
            return jsonify({"error": "不支持的文件格式"}), 400

        # 生成唯一的任务ID
        task_id = str(uuid.uuid4())

        # 创建临时上传目录
        upload_folder = current_app.config.get("UPLOAD_FOLDER", "uploads")
        temp_upload_dir = os.path.join(upload_folder, "temp")
        os.makedirs(temp_upload_dir, exist_ok=True)

        # 保存上传的文件
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{task_id}_{filename}"
        file_path = os.path.join(temp_upload_dir, unique_filename)
        file.save(file_path)

        logger.info(f"公开API文件已保存: {file_path}")

        # 初始化任务状态
        simple_task_status[task_id] = {
            "status": "processing",
            "progress": 0,
            "current_slide": 0,
            "total_slides": 0,
            "file_path": file_path,
            "original_filename": filename,
            "created_at": datetime.now(),
            "error": None,
        }

        # 启动异步翻译任务
        translation_thread = threading.Thread(
            target=execute_simple_translation_task, args=(task_id, file_path, filename)
        )
        translation_thread.daemon = True
        translation_thread.start()

        logger.info(f"公开API翻译任务已启动: {task_id}")

        # 立即返回任务ID
        return jsonify({"task_id": task_id, "status": "started", "message": "翻译任务已启动"})

    except Exception as e:
        logger.error(f"启动公开API翻译任务失败: {str(e)}")
        return jsonify({"error": f"启动翻译任务失败: {str(e)}"}), 500


def execute_simple_translation_task(task_id, file_path, filename):
    """执行简单翻译任务（在后台线程中运行）"""
    try:
        logger.info(f"开始执行公开API翻译任务: {task_id}")

        # 进度回调函数
        def progress_callback(current, total):
            if task_id in simple_task_status:
                progress = int((current / total) * 100) if total > 0 else 0
                simple_task_status[task_id].update(
                    {"progress": progress, "current_slide": current, "total_slides": total}
                )
                logger.info(f"公开API任务 {task_id} 进度: {current}/{total} ({progress}%)")

        # 翻译参数（使用默认值）
        stop_words_list = []
        custom_translations = {}
        select_page = []  # 处理所有页面
        source_language = "en"
        target_language = "zh"
        bilingual_translation = "1"  # 双语模式
        enable_uno_conversion = True  # 默认启用UNO转换

        # 执行翻译
        result = process_presentation(
            file_path,
            stop_words_list,
            custom_translations,
            select_page,
            source_language,
            target_language,
            bilingual_translation,
            progress_callback,
            enable_uno_conversion=enable_uno_conversion,
        )

        if result:
            # 翻译成功
            simple_task_status[task_id].update({"status": "completed", "progress": 100, "completed_at": datetime.now()})
            # 保存翻译后的文件路径
            simple_task_files[task_id] = file_path
            logger.info(f"公开API翻译任务完成: {task_id}")
        else:
            # 翻译失败
            simple_task_status[task_id].update({"status": "failed", "error": "翻译处理失败"})
            logger.error(f"公开API翻译任务失败: {task_id}")

    except Exception as e:
        # 翻译异常
        error_msg = str(e)
        logger.error(f"公开API翻译任务异常: {task_id}, 错误: {error_msg}")
        simple_task_status[task_id].update({"status": "failed", "error": error_msg})


@main.route("/task_status/<task_id>")
def get_simple_task_status(task_id):
    """获取特定任务状态（公开API，不需要认证）"""
    try:
        if task_id not in simple_task_status:
            return jsonify({"status": "not_found", "error": "任务不存在"}), 404

        task = simple_task_status[task_id]

        # 返回任务状态
        response = {
            "status": task["status"],
            "progress": task["progress"],
            "current_slide": task["current_slide"],
            "total_slides": task["total_slides"],
        }

        if task["error"]:
            response["error"] = task["error"]

        return jsonify(response)

    except Exception as e:
        logger.error(f"获取公开API任务状态失败: {str(e)}")
        return jsonify({"status": "error", "error": str(e)}), 500


@main.route("/download/<task_id>")
def download_simple_translated_file(task_id):
    """下载翻译后的文件（公开API，不需要认证）"""
    try:
        if task_id not in simple_task_status:
            return jsonify({"error": "任务不存在"}), 404

        task = simple_task_status[task_id]

        if task["status"] != "completed":
            return jsonify({"error": "任务尚未完成"}), 400

        if task_id not in simple_task_files:
            return jsonify({"error": "翻译文件不存在"}), 404

        file_path = simple_task_files[task_id]

        if not os.path.exists(file_path):
            return jsonify({"error": "文件不存在"}), 404

        return send_file(
            file_path,
            as_attachment=True,
            download_name=f"translated_{task['original_filename']}",
            mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        )

    except Exception as e:
        logger.error(f"下载公开API文件失败: {str(e)}")
        return jsonify({"error": f"下载失败: {str(e)}"}), 500


@main.route("/ppt_translate", methods=["POST"])
def ppt_translate_simple():
    """PPT翻译（公开API，兼容原有接口，不需要认证）"""
    try:
        # 检查是否有文件
        if "file" not in request.files:
            return jsonify({"error": "没有文件"}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "没有选择文件"}), 400

        if not allowed_file(file.filename):
            return jsonify({"error": "不支持的文件格式"}), 400

        # 创建临时上传目录
        upload_folder = current_app.config.get("UPLOAD_FOLDER", "uploads")
        temp_upload_dir = os.path.join(upload_folder, "temp")
        os.makedirs(temp_upload_dir, exist_ok=True)

        # 保存上传的文件
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        file_path = os.path.join(temp_upload_dir, unique_filename)
        file.save(file_path)

        logger.info(f"同步API文件已保存: {file_path}")

        # 翻译参数（使用默认值）
        stop_words_list = []
        custom_translations = {}
        select_page = []  # 处理所有页面
        source_language = "en"
        target_language = "zh"
        bilingual_translation = "1"  # 双语模式
        enable_uno_conversion = True  # 默认启用UNO转换

        # 执行同步翻译
        result = process_presentation(
            file_path,
            stop_words_list,
            custom_translations,
            select_page,
            source_language,
            target_language,
            bilingual_translation,
            enable_uno_conversion=enable_uno_conversion,
        )

        if result:
            logger.info(f"同步API翻译完成: {file_path}")
            # 返回翻译后的文件
            return send_file(
                file_path,
                as_attachment=True,
                download_name=f"translated_{filename}",
                mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            )
        else:
            return jsonify({"error": "翻译处理失败"}), 500

    except Exception as e:
        logger.error(f"同步API翻译失败: {str(e)}")
        return jsonify({"error": f"翻译失败: {str(e)}"}), 500


@main.route("/db_stats")
@login_required
def db_stats():
    """数据库状态页面"""
    if not current_user.is_administrator():
        flash("您没有权限访问此页面")
        return redirect(url_for("main.index"))

    # 获取数据库统计信息
    db_stats = get_db_stats()

    # 获取线程池统计信息
    thread_pool_stats = thread_pool.get_stats()

    # 获取任务队列统计信息
    queue_stats = translation_queue.get_queue_stats()

    return render_template(
        "main/db_stats.html",
        user=current_user,
        db_stats=db_stats,
        thread_pool_stats=thread_pool_stats,
        queue_stats=queue_stats,
    )


@main.route("/db_stats_data")
@login_required
def get_db_stats_data():
    """获取数据库统计数据的API，用于AJAX刷新"""
    if not current_user.is_administrator():
        return jsonify({"error": "没有权限访问此API"}), 403

    # 获取数据库统计信息
    db_stats = get_db_stats()

    return jsonify(db_stats)


@main.route("/recycle_connections", methods=["POST"])
@login_required
def recycle_connections():
    """回收空闲数据库连接"""
    if not current_user.is_administrator():
        return jsonify({"success": False, "message": "没有权限执行此操作"}), 403

    try:
        # 调用翻译队列中的回收连接方法
        result = translation_queue.recycle_idle_connections()

        # 记录操作日志
        logger.info(f"管理员 {current_user.username} 手动回收了数据库空闲连接")

        return jsonify(result)

    except Exception as e:
        logger.error(f"回收数据库连接失败: {str(e)}")
        return jsonify({"success": False, "message": f"回收连接失败: {str(e)}", "error": str(e)}), 500


def get_db_stats():
    """获取数据库连接池统计信息"""
    try:
        engine = db.engine

        # 基本信息
        stats = {
            "engine_name": engine.name,
            "driver_name": engine.driver,
            "url": str(engine.url).replace("://*:*@", "://***:***@"),  # 隐藏敏感信息
            "pool_size": engine.pool.size(),
            "current_size": engine.pool.size(),
            "checked_in": engine.pool.checkedin(),
            "checked_out": engine.pool.checkedout(),
            "overflow": engine.pool.overflow(),
            "max_overflow": engine.pool._max_overflow,
        }

        # 获取连接池配置
        try:
            stats["pool_config"] = {
                "size": engine.pool.size(),
                "max_overflow": engine.pool._max_overflow,
                "timeout": engine.pool._timeout,
                "recycle": engine.pool._recycle,
                "pre_ping": engine.pool._pre_ping,
            }
        except:
            stats["pool_config"] = None

        # 获取已签出连接的详细信息
        checked_out_details = []
        try:
            mutex = engine.pool._mutex
            checked_out = {}

            if hasattr(mutex, "_semlock") and hasattr(engine.pool, "_checked_out"):
                # SQLAlchemy 1.3+
                checked_out = engine.pool._checked_out
            elif hasattr(engine.pool, "_checked_out"):
                # 早期版本
                checked_out = engine.pool._checked_out

            for conn, (ref, traceback, timestamp) in checked_out.items():
                conn_id = str(conn)
                checkout_time = datetime.fromtimestamp(timestamp).strftime("%Y-%m-%d %H:%M:%S")
                duration = time.time() - timestamp
                duration_str = f"{duration:.2f}秒"

                if duration > 3600:
                    hours = int(duration / 3600)
                    minutes = int((duration % 3600) / 60)
                    duration_str = f"{hours}小时{minutes}分钟"
                elif duration > 60:
                    minutes = int(duration / 60)
                    seconds = int(duration % 60)
                    duration_str = f"{minutes}分钟{seconds}秒"

                checked_out_details.append(
                    {
                        "connection_id": conn_id,
                        "checkout_time": checkout_time,
                        "duration": duration_str,
                        "stack_trace": "\n".join(traceback) if traceback else "无堆栈信息",
                    }
                )

            stats["checked_out_details"] = checked_out_details
        except Exception as e:
            stats["checked_out_details"] = []
            logger.warning(f"获取已签出连接详情失败: {str(e)}")

        return stats

    except Exception as e:
        logger.error(f"获取数据库统计信息失败: {str(e)}")
        return {"error": f"获取数据库统计信息失败: {str(e)}"}


@main.route("/system_status", methods=["GET"])
@login_required
def system_status():
    """获取系统状态信息"""
    if not current_user.is_administrator():
        return jsonify({"error": "没有权限访问此API"}), 403

    try:
        # 获取线程池状态
        thread_pool_stats = thread_pool.get_stats()
        thread_pool_health = thread_pool.get_health_status()

        # 获取任务队列状态
        queue_stats = translation_queue.get_queue_stats()

        # 获取数据库连接状态
        db_stats = get_db_stats()

        # 系统内存使用情况
        import psutil

        memory = psutil.virtual_memory()
        memory_stats = {
            "total": memory.total,
            "available": memory.available,
            "used": memory.used,
            "percent": memory.percent,
        }

        # CPU使用情况
        cpu_stats = {
            "percent": psutil.cpu_percent(),
            "count": psutil.cpu_count(),
            "logical_count": psutil.cpu_count(logical=True),
        }

        # 返回汇总状态
        status = {
            "thread_pool": {"stats": thread_pool_stats, "health": thread_pool_health},
            "task_queue": queue_stats,
            "database": db_stats,
            "memory": memory_stats,
            "cpu": cpu_stats,
            "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }

        return jsonify(status)

    except Exception as e:
        logger.error(f"获取系统状态失败: {str(e)}")
        return jsonify({"error": f"获取系统状态失败: {str(e)}"}), 500


@main.route("/system/reset_thread_pool", methods=["POST"])
@login_required
def reset_thread_pool():
    """重置线程池"""
    if not current_user.is_administrator():
        return jsonify({"success": False, "message": "没有权限执行此操作"}), 403

    try:
        # 记录操作日志
        logger.warning(f"管理员 {current_user.username} 正在重置线程池")

        # 获取线程池配置
        stats_before = thread_pool.get_stats()

        # 重新配置线程池
        thread_pool.configure()

        # 获取重置后的状态
        stats_after = thread_pool.get_stats()

        return jsonify({"success": True, "message": "线程池已重置", "before": stats_before, "after": stats_after})

    except Exception as e:
        logger.error(f"重置线程池失败: {str(e)}")
        return jsonify({"success": False, "message": f"重置线程池失败: {str(e)}", "error": str(e)}), 500


@main.route("/system/reset_task_queue", methods=["POST"])
@login_required
def reset_task_queue():
    """重置任务队列"""
    if not current_user.is_administrator():
        return jsonify({"success": False, "message": "没有权限执行此操作"}), 403

    try:
        # 记录操作日志
        logger.warning(f"管理员 {current_user.username} 正在重置任务队列")

        # 获取任务队列状态
        stats_before = translation_queue.get_queue_stats()

        # 停止处理器
        translation_queue.stop_processor()

        # 重新启动处理器
        translation_queue.start_processor()

        # 获取重置后的状态
        stats_after = translation_queue.get_queue_stats()

        return jsonify({"success": True, "message": "任务队列已重置", "before": stats_before, "after": stats_after})

    except Exception as e:
        logger.error(f"重置任务队列失败: {str(e)}")
        return jsonify({"success": False, "message": f"重置任务队列失败: {str(e)}", "error": str(e)}), 500


@main.route("/system_monitoring")
@login_required
def system_monitoring():
    """系统监控页面 - 显示线程池、任务队列和数据库连接状态"""
    # 验证用户是否有管理员权限
    if not current_user.is_administrator:
        flash("您没有访问此页面的权限。", "danger")
        return redirect(url_for("main.index"))

    return render_template("main/system_monitoring.html", user=current_user)


@main.route("/pdf_annotate")
@login_required
def pdf_annotate():
    """PDF注释页面"""
    try:
        # 添加详细的日志
        logger.info("访问 pdf_annotate 页面")
        return render_template("main/pdf_annotate.html")
    except Exception as e:
        logger.error(f"渲染 pdf_annotate 页面出错: {str(e)}")
        # 返回一个简单的错误页面，避免模板渲染问题
        return f"PDF注释功能临时不可用: {str(e)}", 500


import zipfile
import requests
from werkzeug.utils import secure_filename
from datetime import datetime
import os


@main.route('/api/upload_pdf', methods=['POST'])
@login_required
def upload_pdf():
    """上传PDF文件（不开始翻译）"""
    try:
        logger.info("收到PDF文件上传请求")

        # 检查是否有文件上传
        if "file" not in request.files:
            logger.error("未找到上传的文件")
            return jsonify({'success': False, 'error': '未找到上传的文件'}), 400

        original_file = request.files['file']
        if original_file.filename == '':
            logger.error("文件名为空")
            return jsonify({"success": False, "error": "文件名为空"}), 400

        # 生成唯一文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        unique_suffix = uuid.uuid4().hex[:8]
        unique_filename = f"{timestamp}_{unique_suffix}_{secure_filename(original_file.filename)}"
        logger.info(f"生成唯一文件名: {unique_filename}")

        # 获取上传文件夹路径
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        upload_folder = current_app.config['UPLOAD_FOLDER']
        if not os.path.isabs(upload_folder):
            upload_folder = os.path.join(project_root, upload_folder)

        pdf_upload_dir = os.path.join(upload_folder, 'pdf_uploads')
        
        # 确保目录存在
        os.makedirs(pdf_upload_dir, exist_ok=True)

        pdf_path = os.path.join(pdf_upload_dir, unique_filename)
        original_file.save(pdf_path)

        logger.info(f"文件已保存到: {pdf_path}")

        # 验证文件是否正确保存
        if not os.path.exists(pdf_path):
            logger.error(f"文件保存失败，路径不存在: {pdf_path}")
            return jsonify({'success': False, 'error': '文件保存失败'}), 500

        file_size = os.path.getsize(pdf_path)
        logger.info(f"保存的文件大小: {file_size} 字节")

        if file_size == 0:
            logger.error("保存的文件为空")
            return jsonify({'success': False, 'error': '上传的文件为空'}), 400

        return jsonify({
            'success': True,
            'message': 'PDF文件上传成功',
            'file_path': pdf_path,
            'unique_filename': unique_filename,
            'original_filename': original_file.filename
        })

    except Exception as e:
        logger.error(f"上传PDF文件时出错: {e}")
        import traceback
        logger.error(f"错误详情: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': f'上传PDF文件时出错: {str(e)}'}), 500


@main.route('/api/start_pdf_translation', methods=['POST'])
@login_required
def start_pdf_translation():
    """开始PDF翻译（使用已上传的文件）"""
    try:
        logger.info("收到开始PDF翻译请求")

        # 获取已上传的文件信息
        pdf_path = request.form.get('file_path')
        unique_filename = request.form.get('unique_filename')
        original_filename = request.form.get('original_filename')
        
        if not pdf_path or not unique_filename or not original_filename:
            logger.error("缺少必要的文件信息")
            return jsonify({'success': False, 'error': '缺少必要的文件信息'}), 400

        # 验证文件是否存在
        if not os.path.exists(pdf_path):
            logger.error(f"文件不存在: {pdf_path}")
            return jsonify({'success': False, 'error': '文件不存在'}), 400

        # 获取翻译参数
        source_lang = request.form.get('source_lang', 'EN')
        target_lang = request.form.get('target_lang', 'ZH')
        model = request.form.get('model', 'qwen')
        enable_image_ocr = request.form.get('enable_image_ocr', 'false').lower() == 'true'
        
        logger.info(f"PDF翻译参数 - 源语言: {source_lang}, 目标语言: {target_lang}, 模型: {model}, OCR: {enable_image_ocr}")

        # 获取选中的词汇表ID
        selected_vocabulary = request.form.get("selected_vocabulary", "")
        vocabulary_ids = []
        if selected_vocabulary:
            try:
                vocabulary_ids = [int(x.strip()) for x in selected_vocabulary.split(",") if x.strip()]
                logger.info(f"接收到词汇表ID: {vocabulary_ids}")
            except ValueError as e:
                logger.error(f"词汇表ID解析失败: {selected_vocabulary}, 错误: {str(e)}")
                vocabulary_ids = []

        # 构建自定义翻译词典
        custom_translations = {}
        if vocabulary_ids:
            try:
                from app.models import Translation

                translations = Translation.query.filter(
                    Translation.id.in_(vocabulary_ids),
                    db.or_(
                        Translation.user_id == current_user.id,
                        Translation.is_public == True
                    )
                ).all()

                for translation in translations:
                    source_field = {'EN': 'english', 'ZH': 'chinese', 'JA': 'japanese'}.get(source_lang, 'english')
                    target_field = {'EN': 'english', 'ZH': 'chinese', 'JA': 'japanese'}.get(target_lang, 'chinese')

                    source_text = getattr(translation, source_field, '')
                    target_text = getattr(translation, target_field, '')

                    if source_text and target_text:
                        custom_translations[source_text] = target_text

                logger.info(f"构建自定义词典完成，共 {len(custom_translations)} 个词汇")
            except Exception as e:
                logger.error(f"构建自定义词典失败: {str(e)}")
                custom_translations = {}

        # 创建PDF翻译任务
        task_id = str(uuid.uuid4())
        
        # 将任务参数保存到session中，供任务状态查询使用
        session['pdf_task_id'] = task_id
        session['pdf_task_status'] = 'waiting'
        session['pdf_original_filename'] = original_filename
        
        # 初始化任务状态缓存
        with pdf_task_lock:
            pdf_task_status_cache[task_id] = {
                'status': 'processing',
                'message': '正在处理中...'
            }
        
        # 提交任务到线程池（PDF翻译是IO密集型任务）
        from ..utils.thread_pool_executor import thread_pool, TaskType
        task = thread_pool.submit(
            func=process_pdf_translation_async,
            task_type=TaskType.IO_BOUND,
            args=(
                pdf_path,
                original_filename,
                unique_filename,
                source_lang,
                target_lang,
                model,
                enable_image_ocr,
                custom_translations,
                current_user.id,
                task_id
            )
        )

        logger.info(f"PDF翻译任务已提交到队列: {task_id}")
        
        return jsonify({
            'success': True,
            'message': 'PDF翻译任务已创建',
            'task_id': task_id
        })

    except Exception as e:
        logger.error(f"创建PDF翻译任务时出错: {e}")
        import traceback
        logger.error(f"错误详情: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': f'创建PDF翻译任务时出错: {str(e)}'}), 500


@main.route('/translate_pdf', methods=['POST'])
@login_required
def translate_pdf():
    """处理PDF翻译请求（异步模式）"""
    try:
        logger.info("收到PDF翻译请求（异步模式）")

        # 检查是否有文件上传
        if 'file' not in request.files:
            logger.error("未找到上传的文件")
            return jsonify({'success': False, 'error': '未找到上传的文件'}), 400

        original_file = request.files['file']
        if original_file.filename == '':
            logger.error("文件名为空")
            return jsonify({'success': False, 'error': '文件名为空'}), 400

        # 生成唯一文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
        unique_suffix = uuid.uuid4().hex[:8]
        unique_filename = f"{timestamp}_{unique_suffix}_{secure_filename(original_file.filename)}"
        logger.info(f"生成唯一文件名: {unique_filename}")

        # 获取上传文件夹路径
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        upload_folder = current_app.config["UPLOAD_FOLDER"]
        if not os.path.isabs(upload_folder):
            upload_folder = os.path.join(project_root, upload_folder)

        pdf_upload_dir = os.path.join(upload_folder, 'pdf_uploads')
        
        # 确保目录存在
        os.makedirs(pdf_upload_dir, exist_ok=True)

        pdf_path = os.path.join(pdf_upload_dir, unique_filename)
        original_file.save(pdf_path)

        logger.info(f"文件已保存到: {pdf_path}")

        # 验证文件是否正确保存
        if not os.path.exists(pdf_path):
            logger.error(f"文件保存失败，路径不存在: {pdf_path}")
            return jsonify({'success': False, 'error': '文件保存失败'}), 500

        file_size = os.path.getsize(pdf_path)
        logger.info(f"保存的文件大小: {file_size} 字节")

        if file_size == 0:
            logger.error("保存的文件为空")
            return jsonify({'success': False, 'error': '上传的文件为空'}), 400

        # 获取翻译参数
        source_lang = request.form.get('source_lang', 'EN')
        target_lang = request.form.get('target_lang', 'ZH')
        model = request.form.get('model', 'qwen')
        enable_image_ocr = request.form.get('enable_image_ocr', 'false').lower() == 'true'

        # 获取选中的词汇表ID
        selected_vocabulary = request.form.get('selected_vocabulary', '')
        vocabulary_ids = []
        if selected_vocabulary:
            try:
                vocabulary_ids = [int(x.strip()) for x in selected_vocabulary.split(',') if x.strip()]
                logger.info(f"接收到词汇表ID: {vocabulary_ids}")
            except ValueError as e:
                logger.error(f"词汇表ID解析失败: {selected_vocabulary}, 错误: {str(e)}")
                vocabulary_ids = []

        # 构建自定义翻译词典
        custom_translations = {}
        if vocabulary_ids:
            try:
                from app.models import Translation
                translations = Translation.query.filter(
                    Translation.id.in_(vocabulary_ids),
                    db.or_(
                        Translation.user_id == current_user.id,
                        Translation.is_public == True
                    )
                ).all()

                for translation in translations:
                    source_field = {'EN': 'english', 'ZH': 'chinese', 'JA': 'japanese'}.get(source_lang, 'english')
                    target_field = {'EN': 'english', 'ZH': 'chinese', 'JA': 'japanese'}.get(target_lang, 'chinese')

                    source_text = getattr(translation, source_field, '')
                    target_text = getattr(translation, target_field, '')

                    if source_text and target_text:
                        custom_translations[source_text] = target_text

                logger.info(f"构建自定义词典完成，共 {len(custom_translations)} 个词汇")
            except Exception as e:
                logger.error(f"构建自定义词典失败: {str(e)}")
                custom_translations = {}

        # 创建PDF翻译任务
        task_id = str(uuid.uuid4())
        
        # 将任务参数保存到session中，供任务状态查询使用
        session['pdf_task_id'] = task_id
        session['pdf_task_status'] = 'waiting'
        session['pdf_original_filename'] = original_file.filename
        
        # 初始化任务状态缓存
        with pdf_task_lock:
            pdf_task_status_cache[task_id] = {
                'status': 'processing',
                'message': '正在处理中...'
            }
        
        # 提交任务到线程池（PDF翻译是IO密集型任务）
        from ..utils.thread_pool_executor import thread_pool, TaskType
        task = thread_pool.submit(
            func=process_pdf_translation_async,
            task_type=TaskType.IO_BOUND,
            args=(
                pdf_path,
                original_file.filename,
                unique_filename,
                source_lang,
                target_lang,
                model,
                enable_image_ocr,
                custom_translations,
                current_user.id,
                task_id
            )
        )

        logger.info(f"PDF翻译任务已提交到队列: {task_id}")
        
        return jsonify({
            'success': True,
            'message': 'PDF翻译任务已创建',
            'task_id': task_id
        })

    except Exception as e:
        logger.error(f"创建PDF翻译任务时出错: {e}")
        import traceback

        logger.error(f"错误详情: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': f'创建PDF翻译任务时出错: {str(e)}'}), 500

@main.route("/download_translated_pdf/<filename>")
@login_required
def download_translated_pdf(filename):
    """下载翻译后的PDF文件（实际上是Word文档）"""
    try:
        logger.info(f"用户 {current_user.username} 请求下载文件: {filename}")

        from werkzeug.utils import secure_filename

        filename = secure_filename(filename)
        logger.info(f"安全文件名: {filename}")

        project_root = (os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        upload_folder = current_app.config['UPLOAD_FOLDER']
        if not os.path.isabs(upload_folder):
            upload_folder = os.path.join(project_root, upload_folder)

        pdf_output_dir = os.path.join(upload_folder, 'pdf_outputs')
        expected_path = os.path.join(pdf_output_dir, filename)
        absolute_expected_path = os.path.abspath(expected_path)

        logger.info(f"项目根目录: {project_root}")
        logger.info(f"上传文件夹配置: {current_app.config['UPLOAD_FOLDER']}")
        logger.info(f"实际上传文件夹路径: {upload_folder}")
        logger.info(f"PDF输出目录: {pdf_output_dir}")
        logger.info(f"期望的文件路径: {expected_path}")
        logger.info(f"文件绝对路径: {absolute_expected_path}")

        if not os.path.exists(pdf_output_dir):
            logger.error(f"PDF输出目录不存在: {pdf_output_dir}")
            return jsonify({'success': False, 'error': '文件目录不存在'}), 404

        candidate_paths = [absolute_expected_path]
        if not filename.lower().endswith(".docx"):
            candidate_paths.append(os.path.abspath(os.path.join(pdf_output_dir, f"{filename}.docx")))

        file_path = next((path for path in candidate_paths if os.path.exists(path)), None)
        if not file_path:
            logger.error(f"下载文件不存在或不匹配: {absolute_expected_path}")
            logger.info(f"候选路径: {candidate_paths}")
            try:
                logger.info(f"目录现有文件: {os.listdir(pdf_output_dir)}")
            except Exception as list_error:
                logger.warning(f"列出目录失败: {list_error}")
            return jsonify({'success': False, 'error': '文件不存在'}), 404

        download_name = os.path.basename(file_path)

        logger.info(f"准备发送文件给用户: {file_path}")
        logger.info(f"文件大小: {os.path.getsize(file_path)} 字节")
        logger.info(f"文件绝对路径: {os.path.abspath(file_path)}")

        return send_file(file_path, as_attachment=True, download_name=download_name)
    except Exception as e:
        logger.error(f"下载文件时出错: {e}")
        import traceback


@main.route('/api/pdf_translation/delete', methods=['POST'])
@login_required
def delete_pdf_translation():
    """根据文件名删除PDF历史记录及物理文件（若存在记录）"""
    try:
        data = request.get_json(silent=True) or {}
        filename = data.get("filename")
        if not filename:
            return jsonify({"success": False, "error": "缺少文件名"}), 400

        from werkzeug.utils import secure_filename

        filename = secure_filename(filename)

        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        upload_folder = current_app.config["UPLOAD_FOLDER"]
        if not os.path.isabs(upload_folder):
            upload_folder = os.path.join(project_root, upload_folder)
        pdf_output_dir = os.path.join(upload_folder, "pdf_outputs")
        file_path = os.path.join(pdf_output_dir, filename)

        # 删除物理文件
        file_deleted = False
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                file_deleted = True
            except Exception as e:
                logger.warning(f"删除文件失败: {file_path}, {e}")

        # 删除数据库记录（按 stored_filename 匹配）
        from app.models.upload_record import UploadRecord

        record_deleted = False
        try:
            record = UploadRecord.query.filter_by(user_id=current_user.id, stored_filename=filename).first()
            if record:
                db.session.delete(record)
                db.session.commit()
                record_deleted = True
        except Exception as e:
            logger.warning(f"删除数据库记录失败: {e}")
            db.session.rollback()

        return jsonify({"success": True, "file_deleted": file_deleted, "record_deleted": record_deleted})
    except Exception as e:
        logger.error(f"删除PDF历史时出错: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@main.route("/file_management")
@login_required
def file_management():
    """文件管理页面 - 管理员可查看所有用户上传的文件"""
    if not current_user.is_administrator():
        flash('没有权限访问此页面')
        return redirect(url_for('main.index'))

    return render_template('main/file_management.html', user=current_user)


@main.route("/user_management")
@login_required
def user_management():
    """用户管理页面 - 仅管理员可见"""
    if not current_user.is_administrator():
        flash("没有权限访问此页面")
        return redirect(url_for("main.index"))
    return render_template("main/user_management.html", user=current_user)


@main.route("/api/admin/files")
@login_required
def get_admin_files():
    """获取所有用户上传的文件 (仅管理员)"""
    if not current_user.is_administrator():
        return jsonify({'error': '没有权限访问此API'}), 403

    try:
        # 查询所有文件记录
        records = UploadRecord.query.order_by(UploadRecord.upload_time.desc()).all()

        # 构建文件列表，包含用户信息
        files = []
        for record in records:
            # 查询用户信息
            user = User.query.get(record.user_id)
            username = user.username if user else "未知用户"

            # 检查文件是否存在
            file_exists = os.path.exists(os.path.join(record.file_path, record.stored_filename))

            # 使用ISO格式返回时间，让前端正确处理时区
            upload_time = datetime_to_isoformat(record.upload_time)

            files.append({
                'id': record.id,
                'filename': record.filename,
                'stored_filename': record.stored_filename,
                'file_path': record.file_path,
                'file_size': record.file_size,
                'upload_time': upload_time,
                'status': record.status,
                'error_message': record.error_message,
                'user_id': record.user_id,
                'username': username,
                'file_exists': file_exists
            })

        return jsonify({
            'files': files,
            'total': len(files)
        })

    except Exception as e:
        logger.error(f"获取管理员文件列表失败: {str(e)}")
        return jsonify({"error": f"获取文件列表失败: {str(e)}", "files": [], "total": 0}), 500


@main.route("/api/admin/files/<int:record_id>", methods=["DELETE"])
@login_required
def admin_delete_file(record_id):
    """管理员删除文件"""
    if not current_user.is_administrator():
        return jsonify({'error': '没有权限执行此操作'}), 403

    try:
        # 获取上传记录
        record = UploadRecord.query.get_or_404(record_id)

        # 删除物理文件
        file_path = os.path.join(record.file_path, record.stored_filename)
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"管理员删除文件: {file_path}")

        # 删除数据库记录
        db.session.delete(record)
    except Exception as e:
        db.session.rollback()
        logger.error(f"管理员删除文件失败: {str(e)}")
        return jsonify({"success": False, "error": f"删除文件失败: {str(e)}"}), 500
        db.session.commit()
    return jsonify({
        'success': True,
        'message': '文件删除成功'
    })


@main.route("/api/translation_history")
@login_required
def translation_history():
    """获取翻译历史记录"""
    try:
        # 获取查询参数
        file_type = request.args.get('type', '')

        # 构建查询 - 先按用户筛选，不强制状态=completed，避免写库异常导致历史缺失
        query = UploadRecord.query.filter_by(user_id=current_user.id)

        # 按上传时间倒序排列
        records = query.order_by(UploadRecord.upload_time.desc()).all()

        # 格式化记录
        history_records = []
        for record in records:
            # 仅保留PDF翻译生成的记录（目录包含 pdf_outputs）
            try:
                if "pdf_outputs" not in (record.file_path or ""):
                    continue
            except Exception:
                pass
            # 检查文件是否仍然存在
            file_path = os.path.join(record.file_path, record.stored_filename)
            file_exists = os.path.exists(file_path)

            # 如果文件不存在，尝试在pdf_outputs目录中查找
            if not file_exists:
                try:
                    # 构建项目根目录和pdf_outputs路径
                    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                    upload_folder = current_app.config["UPLOAD_FOLDER"]
                    if not os.path.isabs(upload_folder):
                        upload_folder = os.path.join(project_root, upload_folder)
                    pdf_output_dir = os.path.join(upload_folder, 'pdf_outputs')

                    # 在pdf_outputs目录中查找文件
                    potential_file_path = os.path.join(pdf_output_dir, record.stored_filename)
                    if os.path.exists(potential_file_path):
                        file_exists = True
                        file_path = potential_file_path
                        logger.info(f"在pdf_outputs目录中找到历史文件: {file_path}")
                except Exception as e:
                    logger.warning(f"查找历史文件时出错: {e}")

            # 使用ISO格式返回时间，让前端正确处理时区
            upload_time = datetime_to_isoformat(record.upload_time)

            # 直接使用数据库中存储的文件名
            history_records.append(
                {
                    "id": record.id,
                    "filename": record.filename,  # 使用数据库中存储的文件名
                    "stored_filename": getattr(record, "stored_filename", None),
                    "file_size": record.file_size,
                    "upload_time": upload_time,
                    "status": record.status,
                    "file_exists": file_exists,
                }
            )

        # 如果指定了文件类型，则在Python层面进行过滤（避免SQL层面的字段不存在错误）
        if file_type:
            filtered_records = []
            for record in history_records:
                # 由于数据库中可能没有file_type字段，我们只能通过文件名后缀等方式大致判断
                # 这里简化处理，如果需要精确过滤，需要在数据库中添加file_type字段
                if file_type == "pdf_translation":
                    # 简单地通过文件名判断是否为PDF翻译记录
                    if record["filename"].endswith(".docx") or "translated" in record["filename"]:
                        filtered_records.append(record)
                else:
                    # 对于其他类型，暂时不过滤
                    filtered_records.append(record)
            history_records = filtered_records

        return jsonify(history_records)

    except Exception as e:
        logger.error(f"获取PDF翻译历史记录失败: {e}")
        import traceback

        logger.error(f"错误详情: {traceback.format_exc()}")
        return jsonify({"status": "error", "message": "获取历史记录失败"}), 500

        import traceback

        logger.error(f"错误详情: {traceback.format_exc()}")
        return jsonify({"status": "error", "message": "获取历史记录失败"}), 500
        template_path = "./批量上传词汇(模板).xlsx"

        if not os.path.exists(template_path):
            # 如果模板文件不存在，创建它
            create_template_file(template_path)

        return send_file(
            template_path,
            as_attachment=True,
            download_name="批量上传词汇模板.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except Exception as e:
        logger.error(f"下载模板文件失败: {str(e)}")
        return jsonify({"error": f"下载模板文件失败: {str(e)}"}), 500


@main.route("/api/pdf_translation_history")
@login_required
def pdf_translation_history():
    """获取PDF翻译历史记录"""
    try:
        logger.info("[PDF History] 开始查询历史记录")
        # 构建查询 - 只返回状态为 completed 的记录
        query = UploadRecord.query.filter_by(user_id=current_user.id, status='completed')

        # 按上传时间倒序排列
        records = query.order_by(UploadRecord.upload_time.desc()).all()
        logger.info(f"[PDF History] 查询到用户记录数: {len(records)}")

        # 格式化记录
        history_records = []
        for record in records:
            try:
                logger.info(
                    f"[PDF History] 记录: id={record.id}, filename={record.filename}, stored={record.stored_filename}, path={record.file_path}")
            except Exception:
                pass

            # 通过存储文件的后缀来判断是否为PDF翻译记录（PDF翻译结果是.docx文件）
            stored_file_ext = os.path.splitext(record.stored_filename)[1].lower() if record.stored_filename else ''
            if stored_file_ext != '.docx':
                logger.info(
                    f"[PDF History] 过滤非PDF翻译记录: id={record.id}, stored_filename={record.stored_filename}")
                continue

            # 检查文件是否仍然存在
            file_path = os.path.join(record.file_path, record.stored_filename)
            file_exists = os.path.exists(file_path)
            logger.info(f"[PDF History] 文件存在: {file_exists}, full_path={file_path}")

            # 如果文件不存在，尝试在pdf_outputs目录中查找
            if not file_exists:
                try:
                    # 构建项目根目录和pdf_outputs路径
                    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                    upload_folder = current_app.config["UPLOAD_FOLDER"]
                    if not os.path.isabs(upload_folder):
                        upload_folder = os.path.join(project_root, upload_folder)
                    pdf_output_dir = os.path.join(upload_folder, 'pdf_outputs')

                    # 在pdf_outputs目录中查找文件
                    potential_file_path = os.path.join(pdf_output_dir, record.stored_filename)
                    if os.path.exists(potential_file_path):
                        file_exists = True
                        file_path = potential_file_path
                        logger.info(f"在pdf_outputs目录中找到历史文件: {file_path}")
                except Exception as e:
                    logger.warning(f"查找历史文件时出错: {e}")

            # 使用ISO格式返回时间，让前端正确处理时区
            upload_time = datetime_to_isoformat(record.upload_time)

            # 直接使用数据库中存储的文件名
            history_records.append(
                {
                    "id": record.id,
                    "filename": record.filename,  # 使用数据库中存储的文件名
                    "stored_filename": getattr(record, "stored_filename", None),
                    "file_size": record.file_size,
                    "upload_time": upload_time,
                    "status": record.status,
                    "file_exists": file_exists,
                }
            )

        logger.info(f"[PDF History] 返回记录数: {len(history_records)}")
        return jsonify(history_records)

    except Exception as e:
        logger.error(f"获取PDF翻译历史记录失败: {e}")
        import traceback

        logger.error(f"错误详情: {traceback.format_exc()}")
        return jsonify({"status": "error", "message": "获取历史记录失败"}), 500


@main.route("/api/ppt_translation_history")
@login_required
def ppt_translation_history():
    """获取PPT翻译历史记录"""
    try:
        logger.info("[PPT History] 开始查询历史记录")
        # 构建查询 - 只返回状态为 completed 的记录
        query = UploadRecord.query.filter_by(user_id=current_user.id, status="completed")

        # 按上传时间倒序排列
        records = query.order_by(UploadRecord.upload_time.desc()).all()
        logger.info(f"[PPT History] 查询到用户记录数: {len(records)}")

        # 格式化记录
        history_records = []
        for record in records:
            try:
                logger.info(
                    f"[PPT History] 记录: id={record.id}, filename={record.filename}, stored={record.stored_filename}, path={record.file_path}")
            except Exception:
                pass

            # 通过原始文件的后缀来判断是否为PPT翻译记录
            original_file_ext = os.path.splitext(record.filename)[1].lower() if record.filename else ""
            if original_file_ext not in [".ppt", ".pptx"]:
                logger.info(f"[PPT History] 过滤非PPT翻译记录: id={record.id}, filename={record.filename}")
                continue

            # 检查文件是否仍然存在
            file_path = os.path.join(record.file_path, record.stored_filename)
            file_exists = os.path.exists(file_path)
            logger.info(f"[PPT History] 文件存在: {file_exists}, full_path={file_path}")

            # 如果文件不存在，尝试在ppt_outputs目录中查找
            if not file_exists:
                try:
                    # 构建项目根目录和ppt_outputs路径
                    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                    upload_folder = current_app.config["UPLOAD_FOLDER"]
                    if not os.path.isabs(upload_folder):
                        upload_folder = os.path.join(project_root, upload_folder)
                    ppt_output_dir = os.path.join(upload_folder, "ppt_outputs")

                    # 在ppt_outputs目录中查找文件
                    potential_file_path = os.path.join(ppt_output_dir, record.stored_filename)
                    if os.path.exists(potential_file_path):
                        file_exists = True
                        file_path = potential_file_path
                        logger.info(f"在ppt_outputs目录中找到历史文件: {file_path}")
                except Exception as e:
                    logger.warning(f"查找历史文件时出错: {e}")

            # 使用ISO格式返回时间，让前端正确处理时区
            upload_time = datetime_to_isoformat(record.upload_time)

            # 直接使用数据库中存储的文件名
            history_records.append(
                {
                    "id": record.id,
                    "filename": record.filename,  # 使用数据库中存储的文件名
                    "stored_filename": getattr(record, "stored_filename", None),
                    "file_size": record.file_size,
                    "upload_time": upload_time,
                    "status": record.status,
                    "file_exists": file_exists,
                }
            )

        logger.info(f"[PPT History] 返回记录数: {len(history_records)}")
        return jsonify(history_records)

    except Exception as e:
        logger.error(f"获取PPT翻译历史记录失败: {e}")
        import traceback

        logger.error(f"错误详情: {traceback.format_exc()}")
        return jsonify({"status": "error", "message": "获取历史记录失败"}), 500


@main.route("/api/delete_pdf_translation/<int:record_id>", methods=["DELETE"])
@login_required
def delete_pdf_translation_by_id(record_id):
    """删除PDF翻译记录和文件"""
    try:
        # 查询记录，确保是PDF翻译记录
        record = UploadRecord.query.filter_by(
            id=record_id,
            user_id=current_user.id,
            status='completed'
        ).first()

        # 再次确认是PDF翻译记录（通过文件扩展名）
        if record:
            stored_file_ext = os.path.splitext(record.stored_filename)[1].lower() if record.stored_filename else ""
            if stored_file_ext != ".docx":
                record = None

        if not record:
            return jsonify({"status": "error", "message": "记录不存在"}), 404

        # 构建文件路径
        file_path = os.path.join(record.file_path, record.stored_filename)

        # 删除文件（如果存在）
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"已删除PDF翻译文件: {file_path}")
        else:
            # 如果在记录的路径中找不到文件，尝试在pdf_outputs目录中查找
            try:
                project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                upload_folder = current_app.config["UPLOAD_FOLDER"]
                if not os.path.isabs(upload_folder):
                    upload_folder = os.path.join(project_root, upload_folder)
                pdf_output_dir = os.path.join(upload_folder, "pdf_outputs")
                potential_file_path = os.path.join(pdf_output_dir, record.stored_filename)

                if os.path.exists(potential_file_path):
                    os.remove(potential_file_path)
                    logger.info(f"在pdf_outputs目录中删除PDF翻译文件: {potential_file_path}")
            except Exception as e:
                logger.warning(f"在pdf_outputs目录中查找或删除文件时出错: {e}")

        # 删除数据库记录
        db.session.delete(record)
        db.session.commit()

        return jsonify({'status': 'success', 'message': '删除成功'})

    except Exception as e:
        db.session.rollback()
        logger.error(f"删除PDF翻译记录失败: {e}")
        import traceback

        logger.error(f"错误详情: {traceback.format_exc()}")
        return jsonify({'status': 'error', 'message': '删除失败'}), 500


def create_template_file(file_path):
    """创建模板 Excel 文件"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    # 设置表头
    headers = ["english", "chinese", "dutch", "category", "is_public"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = openpyxl.styles.Font(bold=True)
        cell.fill = openpyxl.styles.PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

    # 添加示例数据
    sample_data = [["hello", "你好", "Hallo", "日常；问候", 1], ["sorry", "抱歉", "Pardon", "日常；问候", 0]]

    for row_num, row_data in enumerate(sample_data, 2):
        for col_num, value in enumerate(row_data, 1):
            ws.cell(row=row_num, column=col_num, value=value)

    # 设置列宽度
    column_widths = [20, 20, 20, 30, 10]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    # 保存文件
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    wb.save(file_path)


@main.route('/api/translations/batch_upload', methods=['POST'])
@login_required
def batch_upload_translations():
    """批量上传翻译文件并处理"""
    try:
        if "file" not in request.files:
            return jsonify({"error": "没有文件"}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "没有选择文件"}), 400

        if not allowed_excel_file(file.filename):
            return jsonify({"error": "只支持 Excel 文件 (.xlsx, .xls)"}), 400

        # 获取文件扩展名并验证
        if "." not in file.filename:
            return jsonify({"error": "文件名必须包含扩展名"}), 400

        file_ext = file.filename.rsplit(".", 1)[1].lower()
        if file_ext not in EXCEL_ALLOWED_EXTENSIONS:
            return jsonify(
                {'error': f'不支持的文件格式: .{file_ext}。只支持: {", ".join(EXCEL_ALLOWED_EXTENSIONS)}'}), 400

        # 保存上传的文件
        upload_folder = current_app.config["UPLOAD_FOLDER"]
        user_upload_dir = os.path.join(upload_folder, f"user_{current_user.id}")
        os.makedirs(user_upload_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = secure_filename(file.filename)
        # 确保文件名包含正确的扩展名
        if not filename.lower().endswith(f".{file_ext}"):
            filename = f"{filename}.{file_ext}"

        file_path = os.path.join(user_upload_dir, f"batch_upload_{timestamp}_{filename}")
        file_path = os.path.abspath(file_path)  # 转换为绝对路径
        logger.info(f"文件将保存到: {file_path}")
        logger.info(f"文件扩展名: {file_ext}")

        file.save(file_path)
        logger.info("文件保存成功")

        # 验证文件是否为有效的 Excel 文件
        try:
            import zipfile

            if file_ext == "xlsx":
                # 检查是否为有效的 ZIP 文件（xlsx 实际上是 ZIP 格式）
                with zipfile.ZipFile(file_path, "r") as zip_ref:
                    zip_ref.testzip()
                logger.info("文件是有效的 xlsx 格式")
            elif file_ext == "xls":
                # 对于 xls 文件，检查文件头
                with open(file_path, "rb") as f:
                    header = f.read(8)
                    # Excel 97-2003 的文件头
                    if not header.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
                        raise ValueError("不是有效的 xls 文件")
                logger.info("文件是有效的 xls 格式")
        except Exception as e:
            logger.error(f"文件格式验证失败: {str(e)}")
            os.remove(file_path)  # 删除无效文件
            return jsonify({"error": f"文件格式无效: {str(e)}"}), 400

        # 解析 Excel 文件
        translations_data, errors = parse_excel_file(file_path)

        if errors:
            # 删除临时文件
            os.remove(file_path)
            return jsonify({"error": "文件解析失败", "details": errors[:10]}), 400  # 只返回前10个错误

        if not translations_data:
            # 删除临时文件
            os.remove(file_path)
            return jsonify({"error": "文件中没有有效的翻译数据"}), 400

        # 批量插入数据库
        success_count, error_count, error_details = batch_insert_translations(translations_data, current_user.id)

        # 删除临时文件
        os.remove(file_path)

        result = {
            "message": f"批量上传完成。成功: {success_count}, 失败: {error_count}",
            "success_count": success_count,
            "error_count": error_count,
        }

        if error_details:
            result["errors"] = error_details[:10]  # 只返回前10个错误详情

        return jsonify(result)

    except Exception as e:
        logger.error(f"批量上传翻译失败: {str(e)}")
        logger.error(f"错误类型: {type(e).__name__}")
        import traceback

        logger.error(f"完整错误信息:\n{traceback.format_exc()}")
        return (
            jsonify(
                {
                    "error": f"批量上传失败: {str(e)}",
                    "error_type": type(e).__name__,
                    "file_path": file_path if "file_path" in locals() else None,
                }
            ),
            500,
        )



def parse_excel_file(file_path):
    """解析 Excel 文件，返回翻译数据和错误信息"""
    translations = []
    errors = []

    try:
        logger.info(f"开始解析 Excel 文件: {file_path}")

        # 检查文件是否存在
        if not os.path.exists(file_path):
            errors.append(f"文件不存在: {file_path}")
            return [], errors

        # 检查文件大小
        file_size = os.path.getsize(file_path)
        logger.info(f"文件大小: {file_size} bytes")

        if file_size == 0:
            errors.append("文件为空")
            return [], errors

        logger.info("尝试加载 Excel 文件...")
        wb = openpyxl.load_workbook(file_path, data_only=True)
        logger.info("Excel 文件加载成功")

        ws = wb.active
        logger.info("获取活动工作表成功")

        logger.info(f"工作表名称: {ws.title}")
        logger.info(f"最大行数: {ws.max_row}, 最大列数: {ws.max_column}")

        # 检查表头
        expected_headers = ["english", "chinese", "dutch", "category", "is_public"]
        actual_headers = []

        for col in range(1, len(expected_headers) + 1):
            cell_value = ws.cell(row=1, column=col).value
            if cell_value:
                actual_headers.append(str(cell_value).strip().lower())
            else:
                actual_headers.append("")

        logger.info(f"期望表头: {expected_headers}")
        logger.info(f"实际表头: {actual_headers}")

        if actual_headers != expected_headers:
            errors.append(f"表头不匹配。期望: {expected_headers}, 实际: {actual_headers}")
            return [], errors

        # 解析数据行
        for row_num in range(2, ws.max_row + 1):
            try:
                row_data = {}
                has_data = False

                for col_num, header in enumerate(expected_headers, 1):
                    cell_value = ws.cell(row=row_num, column=col_num).value
                    if cell_value is not None:
                        if isinstance(cell_value, str):
                            cell_value = cell_value.strip()
                        row_data[header] = cell_value
                        if header in ["english", "chinese"] and cell_value:
                            has_data = True
                    else:
                        row_data[header] = None

                # 检查必填字段
                if not row_data.get("english") or not row_data.get("chinese"):
                    if has_data:  # 如果有其他数据但必填字段为空
                        errors.append(f"第{row_num}行: 英文和中文为必填字段")
                    continue

                # 处理 is_public 字段
                if row_data.get("is_public") is not None:
                    if isinstance(row_data["is_public"], str):
                        row_data["is_public"] = row_data["is_public"].lower() in ("1", "true", "yes", "是")
                    elif isinstance(row_data["is_public"], (int, float)):
                        row_data["is_public"] = bool(row_data["is_public"])
                    else:
                        row_data["is_public"] = False
                else:
                    row_data["is_public"] = False

                # 普通用户不能添加公共翻译
                if row_data["is_public"] and not current_user.is_administrator():
                    row_data["is_public"] = False

                translations.append(row_data)

            except Exception as e:
                errors.append(f"第{row_num}行解析失败: {str(e)}")
                continue

    except Exception as e:
        logger.error(f"Excel 文件解析异常: {str(e)}")
        logger.error(f"异常类型: {type(e).__name__}")
        import traceback

        logger.error(f"完整堆栈跟踪:\n{traceback.format_exc()}")
        errors.append(f"文件解析失败: {str(e)}")

    return translations, errors


def batch_insert_translations(translations_data, user_id):
    """批量插入翻译数据到数据库"""
    success_count = 0
    error_count = 0
    error_details = []

    for item in translations_data:
        try:
            # 检查是否已存在相同的翻译
            existing = None
            if item.get("is_public") and current_user.is_administrator():
                # 管理员检查公共翻译
                existing = Translation.query.filter_by(english=item["english"], is_public=True).first()
            else:
                # 普通用户检查自己的私有翻译
                existing = Translation.query.filter_by(user_id=user_id, english=item["english"]).first()

            if existing:
                error_count += 1
                error_details.append(f"英文 '{item['english']}' 已存在")
                continue

            # 创建新的翻译记录
            translation = Translation(
                english=item["english"],
                chinese=item["chinese"],
                dutch=item.get("dutch"),
                category=item.get("category"),
                is_public=item["is_public"],
                user_id=user_id,
            )

            db.session.add(translation)
            success_count += 1

        except Exception as e:
            error_count += 1
            error_details.append(f"插入 '{item.get('english', 'N/A')}' 失败: {str(e)}")
            continue

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        error_details.append(f"数据库提交失败: {str(e)}")
        success_count = 0
        error_count = len(translations_data)

    return success_count, error_count, error_details


EXCEL_ALLOWED_EXTENSIONS = {'xlsx', 'xls'}


def allowed_excel_file(filename):
    if "." not in filename:
        return False
    ext = filename.rsplit(".", 1)[1].lower()
    return ext in EXCEL_ALLOWED_EXTENSIONS
