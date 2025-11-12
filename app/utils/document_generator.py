#!/usr/bin/env python3
"""
æ–‡æ¡£ç”Ÿæˆå™¨æ¨¡å—
ç”¨äºåˆ›å»ºåŒè¯­Wordæ–‡æ¡£ï¼ŒåŒ…å«åŸæ–‡å’Œè¯‘æ–‡çš„å¯¹ç…§æ˜¾ç¤º
"""

import os
import logging
from typing import List, Dict, Tuple
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.enum.style import WD_STYLE_TYPE
import asyncio
from bs4 import BeautifulSoup

# ä¸ºäº†æŒ‰æ®µè½å³æ—¶ç¿»è¯‘ï¼Œå¤ç”¨ç°æœ‰å¼‚æ­¥ç¿»è¯‘èƒ½åŠ›
try:
    from app.function.local_qwen_async import translate_async
except Exception:
    translate_async = None

logger = logging.getLogger(__name__)


class BilingualDocumentGenerator:
    """
    åŒè¯­æ–‡æ¡£ç”Ÿæˆå™¨
    æ”¯æŒåˆ›å»ºåŒ…å«åŸæ–‡å’Œè¯‘æ–‡å¯¹ç…§çš„Wordæ–‡æ¡£
    """
    
    def __init__(self):
        """åˆå§‹åŒ–æ–‡æ¡£ç”Ÿæˆå™¨"""
        self.document = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """è®¾ç½®æ–‡æ¡£æ ·å¼"""
        # è®¾ç½®é»˜è®¤å­—ä½“
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'å®‹ä½“'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
        font.size = Pt(12)
        
        # è®¾ç½®æ ‡é¢˜æ ·å¼
        title_style = self.document.styles['Heading 1']
        title_font = title_style.font
        title_font.name = 'é»‘ä½“'
        title_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        title_font.size = Pt(16)
        title_font.bold = True
        
        # è®¾ç½®äºŒçº§æ ‡é¢˜æ ·å¼
        heading2_style = self.document.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = 'é»‘ä½“'
        heading2_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        heading2_font.size = Pt(14)
        
    def add_heading(self, text: str, level: int = 1) -> None:
        """
        æ·»åŠ æ ‡é¢˜
        
        Args:
            text: æ ‡é¢˜æ–‡æœ¬
            level: æ ‡é¢˜çº§åˆ« (1-6)
        """
        try:
            # æ¸…ç†Markdownç¬¦å·
            cleaned_text = self._strip_inline_markdown(text.strip())
            heading = self.document.add_heading(cleaned_text, level=level)
            # è®¾ç½®ä¸­æ–‡å­—ä½“
            for run in heading.runs:
                run.font.name = 'é»‘ä½“'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
            logger.info(f"æ·»åŠ æ ‡é¢˜: {cleaned_text}")
        except Exception as e:
            logger.error(f"æ·»åŠ æ ‡é¢˜å¤±è´¥: {e}")
            # é™çº§å¤„ç†
            self.add_original_text(text)
    
    def add_original_text(self, text: str) -> None:
        """
        æ·»åŠ åŸæ–‡æ®µè½
        
        Args:
            text: åŸæ–‡æ–‡æœ¬
        """
        if not text.strip():
            self.document.add_paragraph()
            return
            
        try:
            # æ£€æµ‹æ˜¯å¦æ˜¯HTMLè¡¨æ ¼
            if '<table' in text and '</table>' in text:
                # ä½¿ç”¨BeautifulSoupè§£æHTMLè¡¨æ ¼
                soup = BeautifulSoup(text, 'html.parser')
                table = soup.find('table')
                
                if table:
                    # è·å–è¡¨æ ¼çš„è¡Œå’Œåˆ—
                    rows = table.find_all('tr')
                    max_cols = 0
                    cell_contents = []
                    
                    # æå–æ‰€æœ‰å•å…ƒæ ¼å†…å®¹
                    for row in rows:
                        cells = row.find_all(['td', 'th'])
                        max_cols = max(max_cols, len(cells))
                        row_contents = []
                        for cell in cells:
                            # æ¸…ç†å•å…ƒæ ¼æ–‡æœ¬
                            cell_text = self._strip_inline_markdown(cell.get_text().strip())
                            row_contents.append(cell_text)
                        cell_contents.append(row_contents)
                    
                    # åˆ›å»ºåŸæ–‡è¡¨æ ¼
                    word_table = self.document.add_table(rows=len(rows), cols=max_cols)
                    word_table.style = 'Table Grid'
                    
                    # å¡«å……åŸæ–‡è¡¨æ ¼å†…å®¹
                    for i, row_content in enumerate(cell_contents):
                        for j, cell_text in enumerate(row_content):
                            cell = word_table.cell(i, j)
                            cell.text = cell_text
                            # è®¾ç½®å•å…ƒæ ¼æ ¼å¼
                            paragraphs = cell.paragraphs
                            if paragraphs:
                                for paragraph in paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # æ·»åŠ ä¸€ä¸ªç©ºè¡Œåˆ†éš”
                    self.document.add_paragraph()
                    return
                else:
                    # å¦‚æœè¡¨æ ¼è§£æå¤±è´¥ï¼Œé™çº§ä¸ºæ™®é€šæ–‡æœ¬
                    self.document.add_paragraph(text.strip())
                    return
                
            # æ¸…ç†Markdownç¬¦å·
            cleaned_text = self._strip_inline_markdown(text.strip())
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(cleaned_text)
            # åŸæ–‡ä½¿ç”¨é»‘è‰²å­—ä½“
            run.font.color.rgb = RGBColor(0, 0, 0)
            logger.info(f"æ·»åŠ åŸæ–‡: {cleaned_text[:50]}...")
        except Exception as e:
            logger.error(f"æ·»åŠ åŸæ–‡å¤±è´¥: {e}")
            self.document.add_paragraph(text.strip())
    
    def add_translated_text(self, text: str) -> None:
        """
        æ·»åŠ è¯‘æ–‡æ®µè½
        
        Args:
            text: è¯‘æ–‡æ–‡æœ¬
        """
        if not text.strip():
            return
            
        try:
            # æ¸…ç†Markdownç¬¦å·
            cleaned_text = self._strip_inline_markdown(text.strip())
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(cleaned_text)
            # è¯‘æ–‡ä½¿ç”¨ç°è‰²å­—ä½“ä»¥ä¾¿åŒºåˆ†ï¼Œç»Ÿä¸€ä¸ä½¿ç”¨æ–œä½“ï¼Œä¿è¯é£æ ¼ä¸€è‡´
            run.font.color.rgb = RGBColor(96, 96, 96)
            run.italic = False
            
            # è®¾ç½®æ®µè½æ ¼å¼
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(12)
            
            logger.info(f"æ·»åŠ è¯‘æ–‡: {cleaned_text[:50]}...")
        except Exception as e:
            logger.error(f"æ·»åŠ è¯‘æ–‡å¤±è´¥: {e}")
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(text.strip())
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    def add_bilingual_pair(self, original: str, translated: str) -> None:
        """
        æ·»åŠ åŸæ–‡å’Œè¯‘æ–‡å¯¹
        
        Args:
            original: åŸæ–‡
            translated: è¯‘æ–‡
        """
        self.add_original_text(original)
        if translated and not translated.startswith('[ç¿»è¯‘é”™è¯¯:') and translated.strip():
            self.add_translated_text(translated)
        self.document.add_paragraph()  # æ·»åŠ ç©ºè¡Œåˆ†éš”
    
    def add_list_item(self, text: str, numbered: bool = False) -> None:
        """
        æ·»åŠ åˆ—è¡¨é¡¹
        
        Args:
            text: åˆ—è¡¨é¡¹æ–‡æœ¬
            numbered: æ˜¯å¦ä¸ºæœ‰åºåˆ—è¡¨
        """
        try:
            # æ¸…ç†Markdownç¬¦å·
            cleaned_text = self._strip_inline_markdown(text.strip())
            if numbered:
                self.document.add_paragraph(cleaned_text, style='ListNumber')
            else:
                self.document.add_paragraph(cleaned_text, style='ListBullet')
        except Exception as e:
            logger.warning(f"æ·»åŠ åˆ—è¡¨é¡¹å¤±è´¥ï¼Œä½¿ç”¨æ™®é€šæ®µè½: {e}")
            self.add_original_text(text)
    
    def add_bilingual_table(self, html_table: str, source_language: str = "en", target_language: str = "zh", custom_translations: Dict[str, str] = None) -> None:
        """
        å°†HTMLè¡¨æ ¼è½¬æ¢ä¸ºåŒè¯­Wordè¡¨æ ¼å¹¶æ·»åŠ åˆ°æ–‡æ¡£
        ä¼šåˆ›å»ºä¸¤ä¸ªè¡¨æ ¼ï¼šåŸæ–‡è¡¨æ ¼å’Œè¯‘æ–‡è¡¨æ ¼

        Args:
            html_table: HTMLæ ¼å¼çš„è¡¨æ ¼å­—ç¬¦ä¸²
            source_language: æºè¯­è¨€ä»£ç 
            target_language: ç›®æ ‡è¯­è¨€ä»£ç 
            custom_translations: è‡ªå®šä¹‰ç¿»è¯‘å­—å…¸
        """
        try:
            soup = BeautifulSoup(html_table, 'html.parser')
            table = soup.find('table')
            
            if table:
                # è·å–æ‰€æœ‰è¡Œ
                rows = table.find_all('tr')
                if rows:
                    # è®¡ç®—æœ€å¤§åˆ—æ•°å’Œæå–å•å…ƒæ ¼å†…å®¹
                    max_cols = 0
                    cell_contents = []  # åŸæ–‡å†…å®¹
                    translated_contents = []  # è¯‘æ–‡å†…å®¹
                    translation_cache = {}  # ç¿»è¯‘ç¼“å­˜
                    
                    for row in rows:
                        cells = row.find_all(['td', 'th'])
                        max_cols = max(max_cols, len(cells))
                        row_contents = []
                        row_translations = []
                        
                        for cell in cells:
                            # æ¸…ç†å•å…ƒæ ¼æ–‡æœ¬
                            cell_text = self._strip_inline_markdown(cell.get_text().strip())
                            row_contents.append(cell_text)
                            
                            # è·å–ç¿»è¯‘
                            if cell_text in translation_cache:
                                translated_text = translation_cache[cell_text]
                            else:
                                translated_text = _sync_translate_single_text(
                                    cell_text,
                                    source_language=source_language,
                                    target_language=target_language,
                                    custom_translations=custom_translations
                                )
                                translation_cache[cell_text] = translated_text
                            
                            row_translations.append(translated_text if translated_text else cell_text)
                        
                        cell_contents.append(row_contents)
                        translated_contents.append(row_translations)
                    
                    # åˆ›å»ºå¹¶å¡«å……åŸæ–‡è¡¨æ ¼
                    word_table = self.document.add_table(rows=len(rows), cols=max_cols)
                    word_table.style = 'Table Grid'
                    
                    for i, row_content in enumerate(cell_contents):
                        for j, cell_text in enumerate(row_content):
                            cell = word_table.cell(i, j)
                            cell.text = cell_text
                            # è®¾ç½®å•å…ƒæ ¼æ ¼å¼
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # æ·»åŠ ä¸€ä¸ªç©ºè¡Œåˆ†éš”
                    self.document.add_paragraph()
                    
                    # åˆ›å»ºå¹¶å¡«å……è¯‘æ–‡è¡¨æ ¼
                    trans_table = self.document.add_table(rows=len(rows), cols=max_cols)
                    trans_table.style = 'Table Grid'
                    
                    for i, row_translations in enumerate(translated_contents):
                        for j, trans_text in enumerate(row_translations):
                            cell = trans_table.cell(i, j)
                            cell.text = trans_text
                            # è®¾ç½®å•å…ƒæ ¼æ ¼å¼
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(96, 96, 96)
                    
                    logger.info(f"æˆåŠŸæ·»åŠ åŒè¯­HTMLè¡¨æ ¼ï¼Œè¡Œæ•°: {len(rows)}, åˆ—æ•°: {max_cols}")
                    return
            
            # å¦‚æœæ²¡æœ‰æ‰¾åˆ°è¡¨æ ¼ï¼Œé™çº§å¤„ç†ä¸ºæ™®é€šæ–‡æœ¬
            logger.warning("æœªåœ¨HTMLä¸­æ‰¾åˆ°æœ‰æ•ˆè¡¨æ ¼ï¼Œä½œä¸ºæ™®é€šæ–‡æœ¬å¤„ç†")
            self.add_original_text(html_table)
        except Exception as e:
            logger.error(f"æ·»åŠ HTMLè¡¨æ ¼å¤±è´¥: {e}")
            # é™çº§å¤„ç†ï¼Œå°†HTMLä½œä¸ºæ™®é€šæ–‡æœ¬æ·»åŠ 
            self.add_original_text(html_table)

    def add_image(self, image_path: str, width_inches: float = 6.0) -> None:
        """
        æ·»åŠ å›¾ç‰‡åˆ°æ–‡æ¡£
        
        Args:
            image_path: å›¾ç‰‡è·¯å¾„
            width_inches: å›¾ç‰‡å®½åº¦ï¼ˆè‹±å¯¸ï¼‰
        """
        try:
            if os.path.exists(image_path):
                from docx.shared import Inches
                self.document.add_picture(image_path, width=Inches(width_inches))
                logger.info(f"æ·»åŠ å›¾ç‰‡: {image_path}")
            else:
                logger.warning(f"å›¾ç‰‡æ–‡ä»¶ä¸å­˜åœ¨: {image_path}")
        except Exception as e:
            logger.error(f"æ·»åŠ å›¾ç‰‡å¤±è´¥: {e}")
    
    def save(self, file_path: str) -> bool:
        """
        ä¿å­˜æ–‡æ¡£
        
        Args:
            file_path: ä¿å­˜è·¯å¾„
            
        Returns:
            bool: æ˜¯å¦ä¿å­˜æˆåŠŸ
        """
        try:
            # ç¡®ä¿ç›®å½•å­˜åœ¨
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # ä¿å­˜æ–‡æ¡£
            self.document.save(file_path)
            logger.info(f"æ–‡æ¡£ä¿å­˜æˆåŠŸ: {file_path}")
            return True
        except Exception as e:
            logger.error(f"ä¿å­˜æ–‡æ¡£å¤±è´¥: {e}")
            return False

    @staticmethod
    def _strip_inline_markdown(text: str) -> str:
        """
        ç§»é™¤å¸¸è§çš„ Markdown è¡Œå†…è¯­æ³•ï¼Œé¿å…åŸæ ·è¿›å…¥ Wordï¼š
        - ç²—ä½“/æ–œä½“: **text**, __text__, *text*, _text_
        - è¡Œå†…ä»£ç : `code`
        - é“¾æ¥: [text](url) -> text
        - å›¾ç‰‡æ ‡è®°: ![alt](path) -> alt
        - å¼•ç”¨æ®‹ç•™çš„è½¬ä¹‰ç¬¦å·
        """
        try:
            import re
            original_text = text or ""
            s = original_text
            
            # å›¾ç‰‡: ![alt](path) -> alt
            s = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", s)
            # é“¾æ¥: [text](url) -> text
            s = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", s)
            # ç²—ä½“/æ–œä½“åŒ…è£¹: **text** æˆ– __text__ æˆ– *text* æˆ– _text_
            s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
            s = re.sub(r"__([^_]+)__", r"\1", s)
            s = re.sub(r"\*([^*]+)\*", r"\1", s)
            s = re.sub(r"_([^_]+)_", r"\1", s)
            # è¡Œå†…ä»£ç : `code`
            s = re.sub(r"`([^`]+)`", r"\1", s)
            # æ•°å­¦: $...$ / \( ... \) / \[ ... \] -> å»é™¤åŒ…è£¹ç¬¦ï¼Œåªä¿ç•™å†…éƒ¨å†…å®¹
            s = re.sub(r"\$\s*([^$]+?)\s*\$", r"\1", s)
            s = re.sub(r"\\\(\s*([^)]*?)\s*\\\)", r"\1", s)
            s = re.sub(r"\\\[\s*([^]]*?)\s*\\\]", r"\1", s)
            # å¸¸è§ LaTeX è¯­æ³•æ¸…ç†: ^{...} / _{...} å»æ‰æ ‡è®°ï¼Œä¿ç•™å†…å®¹
            s = re.sub(r"\^\s*\{\s*([^}]*)\s*\}", r"\1", s)
            s = re.sub(r"_\s*\{\s*([^}]*)\s*\}", r"\1", s)
            # ç‰¹å®šæ•°å­¦ç¬¦å·æ›¿æ¢: æ”¹è¿›å¯¹ \primeã€\cdotã€\mathsf ç­‰å‘½ä»¤çš„å¤„ç†
            # å¤„ç†å¸¦ç©ºæ ¼çš„å‘½ä»¤ï¼Œå¦‚ { \prime } å’Œ { \cdot }
            s = re.sub(r"\\prime\s*", "â€²", s)  # å°† \prime æ›¿æ¢ä¸º Unicode æ’‡å·
            s = re.sub(r"\{\s*\\prime\s*\}", "â€²", s)  # å°† { \prime } æ›¿æ¢ä¸º Unicode æ’‡å·
            s = re.sub(r"\\cdot\s*", "Â·", s)   # å°† \cdot æ›¿æ¢ä¸º Unicode ä¸­ç‚¹
            s = re.sub(r"\{\s*\\cdot\s*\}", "Â·", s)   # å°† { \cdot } æ›¿æ¢ä¸º Unicode ä¸­ç‚¹
            s = re.sub(r"\\times\s*", "Ã—", s)  # å°† \times æ›¿æ¢ä¸º Unicode ä¹˜å·
            s = re.sub(r"\\leq\s*", "â‰¤", s)    # å°† \leq æ›¿æ¢ä¸º Unicode å°äºç­‰äºå·
            s = re.sub(r"\\geq\s*", "â‰¥", s)    # å°† \geq æ›¿æ¢ä¸º Unicode å¤§äºç­‰äºå·
            s = re.sub(r"\\mathsf\s*\{?\s*([^}]*)\s*\}?", r"\1", s)  # å¤„ç† \mathsf{L} ç­‰æ ¼å¼
            s = re.sub(r"\\mathrm\s*\{?\s*([^}]*)\s*\}?", r"\1", s)  # å¤„ç† \mathrm{R} ç­‰æ ¼å¼
            # å¤„ç†æ›´å¤æ‚çš„å¸¦ç©ºæ ¼èŠ±æ‹¬å·æƒ…å†µ
            s = re.sub(r"\{\s*â€²\s*\}", "â€²", s)   # å°† { â€² } æ›¿æ¢ä¸º â€²
            s = re.sub(r"\{\s*Â·\s*\}", "Â·", s)   # å°† { Â· } æ›¿æ¢ä¸º Â·
            # å¤„ç†å•ç‹¬çš„ ^ å’Œ _ ç¬¦å·
            s = re.sub(r"\^\s*", "", s)  # å»é™¤å•ç‹¬çš„ ^ ç¬¦å·
            s = re.sub(r"_\s*", "", s)   # å»é™¤å•ç‹¬çš„ _ ç¬¦å·
            # ç‰¹åˆ«å¤„ç†ç™¾åˆ†å·å‰çš„åæ–œæ ï¼Œå¦‚30 \%
            s = re.sub(r"\\%", "%", s)
            # å¤„ç†å•†æ ‡ç¬¦å· @ -> Â®
            s = s.replace('@', 'Â®')
            # å»é™¤å¤šä½™çš„åæ–œæ å’ŒèŠ±æ‹¬å·ç©ºæ ¼
            s = s.replace("\\{", "{").replace("\\}", "}").replace("\\ ", " ")
            # è§„èŒƒè¿å­—ç¬¦ä¸ç©ºæ ¼: é¿å… ' - ' æ®‹ç•™ç©ºæ ¼
            s = re.sub(r"\s*-\s*", "-", s)
            # æŠ˜å å¤šç©ºæ ¼
            s = re.sub(r"\s+", " ", s).strip()
            # å‰©ä½™çš„è½¬ä¹‰åæ–œæ 
            s = s.replace("\\*", "*").replace("\\_", "_").replace("\\#", "#").replace("\\`", "`")
            
            # è®°å½•æ–‡æœ¬æ¸…ç†æ“ä½œ
            if s != original_text:
                logger.debug(f"æ–‡æœ¬æ¸…ç†å®Œæˆ: '{original_text[:30]}...' -> '{s[:30]}...'")
            
            return s
        except Exception as e:
            logger.error(f"æ–‡æœ¬æ¸…ç†è¿‡ç¨‹ä¸­å‡ºé”™: {e}")
            return text or ""


def create_bilingual_word_document(
    content_lines: List[str], 
    translation_dict: Dict[str, str],
    output_path: str
) -> bool:
    """
    åˆ›å»ºåŒè¯­Wordæ–‡æ¡£
    
    Args:
        content_lines: åŸå§‹å†…å®¹è¡Œåˆ—è¡¨
        translation_dict: ç¿»è¯‘å­—å…¸ {åŸæ–‡: è¯‘æ–‡}
        output_path: è¾“å‡ºæ–‡ä»¶è·¯å¾„
        
    Returns:
        bool: æ˜¯å¦åˆ›å»ºæˆåŠŸ
    """
    try:
        generator = BilingualDocumentGenerator()
        
        i = 0
        while i < len(content_lines):
            line = content_lines[i].strip()
            
            if not line:
                generator.document.add_paragraph()
                i += 1
                continue
            
            # å¤„ç†æ ‡é¢˜ (# å¼€å¤´)
            if line.startswith('#'):
                # è®¡ç®—æ ‡é¢˜çº§åˆ«
                level = 1
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                level = min(level, 6)  # æœ€å¤šæ”¯æŒ6çº§æ ‡é¢˜
                
                # æå–æ ‡é¢˜æ–‡æœ¬
                title_text = line.lstrip('# ').strip()
                generator.add_heading(title_text, level)
                
                # æ£€æŸ¥ä¸‹ä¸€è¡Œæ˜¯å¦ä¸ºè¯‘æ–‡
                if i + 1 < len(content_lines):
                    next_line = content_lines[i + 1].strip()
                    if next_line.startswith('ã€è¯‘æ–‡ã€‘'):
                        translated_text = next_line[5:].strip()  # å»æ‰"ã€è¯‘æ–‡ã€‘"å‰ç¼€
                        generator.add_translated_text(translated_text)
                        i += 2  # è·³è¿‡æ ‡é¢˜å’Œè¯‘æ–‡è¡Œ
                        continue
                
                i += 1
                continue
            
            # å¤„ç†æ— åºåˆ—è¡¨ (* æˆ– - å¼€å¤´)
            if line.startswith('* ') or line.startswith('- '):
                list_text = line[2:].strip()
                generator.add_list_item(list_text, numbered=False)
                
                # æ£€æŸ¥ä¸‹ä¸€è¡Œæ˜¯å¦ä¸ºè¯‘æ–‡
                if i + 1 < len(content_lines):
                    next_line = content_lines[i + 1].strip()
                    if next_line.startswith('ã€è¯‘æ–‡ã€‘'):
                        translated_text = next_line[5:].strip()
                        generator.add_translated_text(translated_text)
                        i += 2  # è·³è¿‡åˆ—è¡¨é¡¹å’Œè¯‘æ–‡è¡Œ
                        continue
                
                i += 1
                continue
            
            # å¤„ç†æœ‰åºåˆ—è¡¨ (æ•°å­—. å¼€å¤´)
            if line.startswith(tuple(f"{n}. " for n in range(1, 100))):  # æ”¯æŒ1-99çš„ç¼–å·
                dot_index = line.find('. ')
                if dot_index > 0:
                    list_text = line[dot_index + 2:].strip()
                    generator.add_list_item(list_text, numbered=True)
                    
                    # æ£€æŸ¥ä¸‹ä¸€è¡Œæ˜¯å¦ä¸ºè¯‘æ–‡
                    if i + 1 < len(content_lines):
                        next_line = content_lines[i + 1].strip()
                        if next_line.startswith('ã€è¯‘æ–‡ã€‘'):
                            translated_text = next_line[5:].strip()
                            generator.add_translated_text(translated_text)
                            i += 2  # è·³è¿‡åˆ—è¡¨é¡¹å’Œè¯‘æ–‡è¡Œ
                            continue
                    
                    i += 1
                    continue
            
            # å¤„ç†è¯‘æ–‡æ ‡è®°è¡Œ
            if line.startswith('ã€è¯‘æ–‡ã€‘'):
                # è¿™ç§æƒ…å†µç†è®ºä¸Šä¸åº”è¯¥å‡ºç°ï¼Œå› ä¸ºæˆ‘ä»¬ä¼šåœ¨æ·»åŠ åŸæ–‡æ—¶å¤„ç†å¯¹åº”çš„è¯‘æ–‡
                translated_text = line[5:].strip()
                generator.add_translated_text(translated_text)
                i += 1
                continue
            
            # å¤„ç†æ™®é€šæ®µè½
            original_text = line
            translated_text = ""
            
            # å°è¯•æŸ¥æ‰¾ç¿»è¯‘
            if original_text in translation_dict:
                translated_text = translation_dict[original_text]
            else:
                # ä½¿ç”¨æˆ‘ä»¬ä¹‹å‰å®ç°çš„æ¨¡ç³ŠåŒ¹é…é€»è¾‘
                for key, value in translation_dict.items():
                    if (original_text.strip() in key.strip() or 
                        key.strip() in original_text.strip()):
                        translated_text = value
                        break
            
            generator.add_bilingual_pair(original_text, translated_text)
            i += 1
        
        # ä¿å­˜æ–‡æ¡£
        return generator.save(output_path)
        
    except Exception as e:
        logger.error(f"åˆ›å»ºåŒè¯­Wordæ–‡æ¡£å¤±è´¥: {e}")
        import traceback
        logger.error(f"é”™è¯¯è¯¦æƒ…: {traceback.format_exc()}")
        return False


def process_markdown_to_bilingual_doc(
    markdown_content: str,
    translation_dict: Dict[str, str],
    output_path: str
) -> bool:
    """
    å°†Markdownå†…å®¹å¤„ç†ä¸ºåŒè¯­Wordæ–‡æ¡£
    
    Args:
        markdown_content: Markdownæ ¼å¼çš„å†…å®¹
        translation_dict: ç¿»è¯‘å­—å…¸ {åŸæ–‡: è¯‘æ–‡}
        output_path: è¾“å‡ºæ–‡ä»¶è·¯å¾„
        
    Returns:
        bool: æ˜¯å¦å¤„ç†æˆåŠŸ
    """
    try:
        # æŒ‰è¡Œåˆ†å‰²å†…å®¹
        content_lines = markdown_content.split('\n')
        
        # åˆ›å»ºåŒè¯­æ–‡æ¡£
        success = create_bilingual_word_document(content_lines, translation_dict, output_path)
        
        if success:
            logger.info(f"æˆåŠŸåˆ›å»ºåŒè¯­Wordæ–‡æ¡£: {output_path}")
        else:
            logger.error("åˆ›å»ºåŒè¯­Wordæ–‡æ¡£å¤±è´¥")
            
        return success
        
    except Exception as e:
        logger.error(f"å¤„ç†Markdownåˆ°åŒè¯­æ–‡æ¡£å¤±è´¥: {e}")
        import traceback
        logger.error(f"é”™è¯¯è¯¦æƒ…: {traceback.format_exc()}")
        return False


def _sync_translate_single_text(text: str,
                                source_language: str = "en",
                                target_language: str = "zh",
                                custom_translations: Dict[str, str] = None) -> str:
    """
    ä½¿ç”¨ç°æœ‰çš„å¼‚æ­¥ç¿»è¯‘æ¥å£å¯¹å•æ¡æ–‡æœ¬è¿›è¡ŒåŒæ­¥ç¿»è¯‘ï¼Œè¿”å›è¯‘æ–‡ã€‚

    è‹¥ç¿»è¯‘ä¸å¯ç”¨æˆ–å¤±è´¥ï¼Œè¿”å›ç©ºå­—ç¬¦ä¸²ä»¥ä¾¿è°ƒç”¨æ–¹é™çº§å¤„ç†ã€‚
    """
    if not text or not text.strip():
        return ""
    if translate_async is None:
        return ""
    try:
        async def _run_once() -> str:
            mapping = await translate_async(
                text.strip(),
                field="é€šç”¨",  # PDFç¿»è¯‘ä½¿ç”¨é€šç”¨é¢†åŸŸï¼Œè·³è¿‡é¢†åŸŸæ£€æµ‹
                stop_words=[],
                custom_translations=custom_translations or {},
                source_language=source_language,
                target_language=target_language
            )
            # translate_async è¿”å› {åŸæ–‡: è¯‘æ–‡}
            if isinstance(mapping, dict):
                # ä¼˜å…ˆç²¾ç¡®åŒ¹é…æ•´è¡Œ
                key = text.strip()
                if key in mapping:
                    return mapping[key] or ""
                # æ— ç²¾ç¡®åŒ¹é…æ—¶ï¼Œåˆå¹¶æ‰€æœ‰è¯‘æ–‡ï¼Œé¿å…åªå–é¦–å¥
                merged_values: list[str] = []
                for _, v in mapping.items():
                    if isinstance(v, str) and v.strip():
                        merged_values.append(v.strip())
                if merged_values:
                    return " ".join(merged_values)
            # é™çº§é‡è¯•ï¼šä½¿ç”¨å¤‡ç”¨ç¿»è¯‘å™¨ç›´è¯‘æ•´æ®µ
            try:
                from app.function.image_ocr.translator import QwenTranslator
                # å°†å†…éƒ¨è¯­è¨€ä»£ç æ˜ å°„ä¸ºäººç±»å¯è¯»ï¼ˆå¤‡ç”¨ç¿»è¯‘å™¨æç¤ºè¯ç”¨ï¼‰
                def _map_lang_name(code: str) -> str:
                    c = (code or "").lower()
                    if c.startswith("zh") or c == "cn" or c == "chinese":
                        return "ä¸­æ–‡"
                    if c.startswith("en") or c == "english":
                        return "è‹±æ–‡"
                    if c.startswith("ja") or c == "japanese":
                        return "æ—¥æ–‡"
                    return "è‹±æ–‡" if c else "è‹±æ–‡"
                qt = QwenTranslator(target_language=_map_lang_name(target_language))
                fallback = qt.translate_text(text.strip(), source_language=_map_lang_name(source_language))
                return fallback or ""
            except Exception:
                return ""

        # ç‹¬ç«‹è¿è¡Œäº‹ä»¶å¾ªç¯
        try:
            return asyncio.run(_run_once())
        except RuntimeError:
            # è‹¥å·²æœ‰äº‹ä»¶å¾ªç¯ï¼ˆå°‘è§äºFlaskï¼‰ï¼Œåˆ™åˆ›å»ºæ–°å¾ªç¯æ‰§è¡Œ
            loop = asyncio.new_event_loop()
            try:
                return loop.run_until_complete(_run_once())
            finally:
                loop.close()
    except Exception:
        return ""


def translate_markdown_to_bilingual_doc(
    markdown_content: str,
    output_path: str,
    source_language: str = "en",
    target_language: str = "zh",
    image_base_dir: str | None = None,
    custom_translations: Dict[str, str] = None,
    image_ocr_results: list[dict] | None = None
) -> bool:
    """
    å°†Markdownå†…å®¹æŒ‰"æ ‡é¢˜/æ®µè½ â†’ é€æ¡ç¿»è¯‘ â†’ ç«‹å³å†™å…¥Word(åŸæ–‡åœ¨å‰ï¼Œè¯‘æ–‡åœ¨å)"çš„æ–¹å¼ç”ŸæˆåŒè¯­Wordæ–‡æ¡£ã€‚

    - æ ‡é¢˜(`#`çº§åˆ«)å†™å…¥ä¸ºå¯¹åº”Headingï¼Œç„¶åç´§è·Ÿè¯‘æ–‡æ®µè½
    - æ— åº/æœ‰åºåˆ—è¡¨é¡¹å†™å…¥åˆ—è¡¨é¡¹ï¼Œç„¶åç´§è·Ÿè¯‘æ–‡æ®µè½
    - æ™®é€šæ®µè½ä½¿ç”¨ add_bilingual_pairï¼ŒåŸæ–‡åç´§è·Ÿè¯‘æ–‡
    - ç©ºè¡Œä¿æŒ
    
    å‚æ•°:
        markdown_content: Markdownå†…å®¹
        output_path: è¾“å‡ºWordæ–‡æ¡£è·¯å¾„
        source_language: æºè¯­è¨€
        target_language: ç›®æ ‡è¯­è¨€
        image_base_dir: å›¾ç‰‡åŸºç¡€ç›®å½•
        custom_translations: è‡ªå®šä¹‰ç¿»è¯‘å­—å…¸
        image_ocr_results: å›¾ç‰‡OCRè¯†åˆ«ç»“æœåˆ—è¡¨ï¼Œæ¯ä¸ªå…ƒç´ åŒ…å«ï¼š
            - success: bool - OCRæ˜¯å¦æˆåŠŸ
            - image_path: str - å›¾ç‰‡è·¯å¾„
            - ocr_text_combined: str - OCRè¯†åˆ«çš„æ–‡æœ¬
            - translation_text_combined: str - ç¿»è¯‘åçš„æ–‡æœ¬

    è‹¥ç¿»è¯‘ä¸å¯ç”¨æˆ–å¤±è´¥ï¼Œä¾ç„¶å†™å…¥åŸæ–‡ï¼Œè¯‘æ–‡ç•™ç©ºã€‚
    """
    try:
        # åˆ›å»ºå›¾ç‰‡è·¯å¾„åˆ°OCRç»“æœçš„æ˜ å°„
        ocr_results_map = {}
        if image_ocr_results:
            import os as _os
            for result in image_ocr_results:
                if result.get('success'):
                    img_path = result.get('image_path', '')
                    if img_path:
                        # ä½¿ç”¨è§„èŒƒåŒ–è·¯å¾„ä½œä¸ºkey
                        normalized_path = _os.path.normpath(img_path)
                        ocr_results_map[normalized_path] = result
            if ocr_results_map:
                logger.info(f"ğŸ“Š å·²åŠ è½½ {len(ocr_results_map)} ä¸ªå›¾ç‰‡çš„OCRç»“æœåˆ°æ˜ å°„å­—å…¸")
                logger.info("OCRç»“æœæ˜ å°„ä¸­çš„è·¯å¾„:")
                for idx, (path, result) in enumerate(list(ocr_results_map.items())[:3], 1):
                    logger.info(f"  {idx}. {path}")
                    logger.info(f"     OCRæ–‡æœ¬é•¿åº¦: {len(result.get('ocr_text_combined', ''))}")
                    logger.info(f"     ç¿»è¯‘æ–‡æœ¬é•¿åº¦: {len(result.get('translation_text_combined', ''))}")
                if len(ocr_results_map) > 3:
                    logger.info(f"  ... è¿˜æœ‰ {len(ocr_results_map) - 3} ä¸ªå›¾ç‰‡")
            else:
                logger.warning("âš ï¸ OCRç»“æœæ˜ å°„å­—å…¸ä¸ºç©º - æœªæ·»åŠ ä»»ä½•å›¾ç‰‡")
        else:
            logger.info("â„¹ï¸ image_ocr_resultså‚æ•°ä¸ºNoneæˆ–ç©ºåˆ—è¡¨ï¼Œè·³è¿‡OCRç»“æœæ˜ å°„")
        
        generator = BilingualDocumentGenerator()
        if not markdown_content:
            return generator.save(output_path)

        lines = markdown_content.split('\n')

        # ç®€å•çš„ä¸­æ–‡æ£€æµ‹(ç”¨äºè·³è¿‡å·²æ˜¯ä¸­æ–‡çš„æ–‡æœ¬)
        def is_mostly_chinese(s: str) -> bool:
            if not s:
                return False
            total = len(s)
            if total == 0:
                return False
            import re
            cjk = len(re.findall(r"[\u4e00-\u9fff]", s))
            return (cjk / total) > 0.5

        # è§£æ markdown ä¸ºå—ï¼šæ ‡é¢˜ã€åˆ—è¡¨é¡¹ã€å›¾ç‰‡ã€æ™®é€šæ®µè½(åˆå¹¶è¿ç»­è¡Œ)
        blocks: list[dict] = []
        i = 0
        while i < len(lines):
            raw_line = lines[i]
            line = (raw_line or "").rstrip('\r')

            # è·³è¿‡å¸¦æœ‰å›ºå®šè¯‘æ–‡æ ‡è®°çš„è¡Œ
            if line.strip().startswith('ã€è¯‘æ–‡ã€‘'):
                i += 1
                continue

            # å›¾ç‰‡è¡Œ(åªå¤„ç†æ•´è¡Œå›¾ç‰‡è¯­æ³•): ![alt](path)
            stripped = line.lstrip()
            if stripped.startswith('![') and '](' in stripped and stripped.endswith(')'):
                try:
                    alt_end = stripped.find('](')
                    path = stripped[alt_end+2:-1].strip()
                    blocks.append({'type': 'image', 'path': path, 'alt': stripped[2:alt_end]})
                    i += 1
                    continue
                except Exception:
                    pass

            # æ ‡é¢˜
            if stripped.startswith('#'):
                level = 0
                for ch in stripped:
                    if ch == '#':
                        level += 1
                    else:
                        break
                level = max(1, min(level, 6))
                title_text = stripped[level:].strip()
                blocks.append({'type': 'heading', 'level': level, 'text': title_text})
                i += 1
                continue

            # åˆ—è¡¨é¡¹
            if stripped.startswith('* ') or stripped.startswith('- '):
                item_text = stripped[2:].strip()
                blocks.append({'type': 'ul_item', 'text': item_text})
                i += 1
                continue

            dot_index = stripped.find('. ')
            if dot_index > 0 and stripped[:dot_index].isdigit():
                item_text = stripped[dot_index+2:].strip()
                blocks.append({'type': 'ol_item', 'text': item_text})
                i += 1
                continue

            # ç©ºè¡Œ -> ä½œä¸ºæ®µè½åˆ†éš”
            if not line.strip():
                blocks.append({'type': 'blank'})
                i += 1
                continue

            # åˆå¹¶è¿ç»­çš„éç©ºæ™®é€šæ–‡æœ¬è¡Œä¸ºä¸€ä¸ªæ®µè½
            paragraph_lines = [line.strip()]
            j = i + 1
            while j < len(lines):
                nxt = (lines[j] or "").rstrip('\r')
                nxt_stripped = nxt.lstrip()
                if not nxt.strip():
                    break
                if nxt_stripped.startswith('#') or nxt_stripped.startswith('![') or \
                   nxt_stripped.startswith('* ') or nxt_stripped.startswith('- '):
                    break
                di = nxt_stripped.find('. ')
                if di > 0 and nxt_stripped[:di].isdigit():
                    break
                if nxt_stripped.startswith('ã€è¯‘æ–‡ã€‘'):
                    break
                paragraph_lines.append(nxt.strip())
                j += 1
            merged = ' '.join(paragraph_lines).strip()
            if merged:
                blocks.append({'type': 'paragraph', 'text': merged})
            i = j if j > i else (i + 1)

        # å†™å…¥åˆ° Wordï¼Œä½¿ç”¨å»é‡ç¼“å­˜é¿å…é‡å¤
        cache: dict[str, str] = {}
        for blk in blocks:
            btype = blk.get('type')
            if btype == 'blank':
                generator.document.add_paragraph()
                continue

            if btype == 'image':
                path = blk.get('path') or ''
                if path:
                    # å¦‚æœæ˜¯ç›¸å¯¹è·¯å¾„ï¼Œåˆ™åŸºäº image_base_dir è§£æ
                    final_path = path
                    if image_base_dir and not (path.startswith('http://') or path.startswith('https://') or 
                                               path.startswith('/') or (len(path) > 1 and path[1] == ':')):
                        import os as _os
                        final_path = _os.path.join(image_base_dir, path)
                    
                    # å°è¯•æ’å…¥å›¾ç‰‡
                    image_inserted = False
                    try:
                        logger.info(f"å°è¯•æ’å…¥å›¾ç‰‡: {final_path}")
                        if _os.path.exists(final_path):
                            generator.document.add_picture(final_path, width=Inches(6.0))
                            image_inserted = True
                            logger.info(f"  æˆåŠŸæ’å…¥å›¾ç‰‡: {final_path}")
                        else:
                            logger.warning(f"âš ï¸ å›¾ç‰‡æ–‡ä»¶ä¸å­˜åœ¨: {final_path}")
                    except Exception as img_error:
                        logger.error(f"âŒ æ’å…¥å›¾ç‰‡å¤±è´¥: {final_path}, é”™è¯¯: {img_error}")
                    
                    # æ£€æŸ¥æ˜¯å¦æœ‰OCRç»“æœå¹¶æ·»åŠ åˆ°æ–‡æ¡£
                    if ocr_results_map and image_inserted:
                        normalized_path = _os.path.normpath(final_path)
                        logger.info(f"ğŸ” æŸ¥æ‰¾å›¾ç‰‡OCRç»“æœ...")
                        logger.info(f"   åŸå§‹è·¯å¾„: {final_path}")
                        logger.info(f"   è§„èŒƒåŒ–è·¯å¾„: {normalized_path}")
                        logger.info(f"   æ˜ å°„ä¸­æœ‰ {len(ocr_results_map)} ä¸ªå›¾ç‰‡")
                        ocr_result = ocr_results_map.get(normalized_path)
                        
                        if ocr_result:
                            logger.info(f"  æ‰¾åˆ°OCRç»“æœï¼")
                            ocr_text = ocr_result.get('ocr_text_combined', '').strip()
                            translation_text = ocr_result.get('translation_text_combined', '').strip()
                            logger.info(f"   OCRæ–‡æœ¬: {len(ocr_text)} å­—ç¬¦")
                            logger.info(f"   ç¿»è¯‘æ–‡æœ¬: {len(translation_text)} å­—ç¬¦")
                            
                            if ocr_text or translation_text:
                                # æ·»åŠ OCRç»“æœçš„æ ‡é¢˜
                                generator.document.add_paragraph(
                                    f"ã€å›¾ç‰‡æ–‡å­—è¯†åˆ«ã€‘",
                                    style='Heading 3'
                                )
                                
                                # æ·»åŠ OCRåŸæ–‡
                                if ocr_text:
                                    ocr_para = generator.document.add_paragraph()
                                    ocr_run = ocr_para.add_run(f"åŸæ–‡: {ocr_text}")
                                    ocr_run.font.size = Pt(10)
                                    ocr_run.font.color.rgb = RGBColor(70, 70, 70)
                                    logger.info(f"    å·²æ·»åŠ OCRåŸæ–‡åˆ°Word ({len(ocr_text)} å­—ç¬¦)")
                                
                                # æ·»åŠ ç¿»è¯‘æ–‡æœ¬
                                if translation_text:
                                    trans_para = generator.document.add_paragraph()
                                    trans_run = trans_para.add_run(f"è¯‘æ–‡: {translation_text}")
                                    trans_run.font.size = Pt(10)
                                    trans_run.font.color.rgb = RGBColor(0, 102, 204)
                                    logger.info(f"    å·²æ·»åŠ OCRè¯‘æ–‡åˆ°Word ({len(translation_text)} å­—ç¬¦)")
                                
                                # æ·»åŠ åˆ†éš”çº¿
                                generator.document.add_paragraph("â”€" * 50)
                                logger.info(f"    OCRç»“æœå·²å®Œæ•´æ·»åŠ åˆ°Wordæ–‡æ¡£")
                            else:
                                logger.warning(f"  âš ï¸ OCRæ–‡æœ¬å’Œç¿»è¯‘éƒ½ä¸ºç©º")
                        else:
                            logger.warning(f"âŒ æœªæ‰¾åˆ°åŒ¹é…çš„OCRç»“æœ")
                            logger.warning(f"   å›¾ç‰‡æ–‡ä»¶å: {_os.path.basename(final_path)}")
                            logger.warning(f"   è§„èŒƒåŒ–è·¯å¾„: {normalized_path}")
                            if ocr_results_map:
                                logger.warning(f"   æ˜ å°„ä¸­çš„è·¯å¾„ç¤ºä¾‹ï¼ˆå‰3ä¸ªï¼‰:")
                                for idx, map_path in enumerate(list(ocr_results_map.keys())[:3], 1):
                                    logger.warning(f"     {idx}. {map_path}")
                    elif not image_inserted:
                        logger.warning(f"âš ï¸ å›¾ç‰‡æœªæˆåŠŸæ’å…¥ï¼Œè·³è¿‡OCRç»“æœå¤„ç†")
                    elif not ocr_results_map:
                        logger.info(f"â„¹ï¸ æ²¡æœ‰OCRç»“æœæ˜ å°„ï¼Œè·³è¿‡OCRæ–‡æœ¬æ·»åŠ ")
                continue

            text = blk.get('text', '').strip()
            if not text:
                continue

            # ä»…åœ¨è‹±ç¿»ä¸­åœºæ™¯ä¸‹è·³è¿‡å·²æ˜¯ä¸­æ–‡çš„åŸæ–‡ï¼Œé¿å…é‡å¤è¯‘æ–‡
            # å½“ç›®æ ‡è¯­è¨€æ˜¯ä¸­æ–‡æ—¶ï¼Œæ£€æµ‹åŸæ–‡æ˜¯å¦ä¸»è¦æ˜¯ä¸­æ–‡ï¼Œå¦‚æœæ˜¯åˆ™è·³è¿‡ç¿»è¯‘
            # æ³¨æ„ï¼šä¸è¦è·³è¿‡ä¸­æ–‡å†…å®¹ï¼Œå³ä½¿æ˜¯è‹±ç¿»ä¸­åœºæ™¯ï¼Œå› ä¸ºå¯èƒ½ç”¨æˆ·éœ€è¦ä¸­è‹±å¯¹ç…§
            # æ³¨é‡Šæ‰è·³è¿‡é€»è¾‘ï¼Œç¡®ä¿æ‰€æœ‰å†…å®¹éƒ½èƒ½è¢«ç¿»è¯‘å¹¶æ·»åŠ åˆ°æ–‡æ¡£ä¸­
            # if target_language.lower().startswith('zh') and is_mostly_chinese(text):
            #     generator.add_original_text(text)
            #     continue

            if btype == 'heading':
                level = int(blk.get('level', 1))
                generator.add_heading(text, level)

                # æ£€æŸ¥æ–‡æœ¬è¯­è¨€ç‰¹æ€§
                is_chinese_content = is_mostly_chinese(text)
                is_target_english = target_language.lower().startswith('en')
                is_target_chinese = target_language.lower().startswith('zh')

                # æ£€æŸ¥æ˜¯å¦æ˜¯è‹±æ–‡å†…å®¹ï¼ˆç®€å•æ£€æµ‹ï¼šåŒ…å«è¾ƒå¤šæ‹‰ä¸å­—æ¯ï¼‰
                def is_mostly_english(s: str) -> bool:
                    if not s:
                        return False
                    total = len(s)
                    if total == 0:
                        return False
                    import re
                    latin = len(re.findall(r"[a-zA-Z]", s))
                    return (latin / total) > 0.5

                is_english_content = is_mostly_english(text)

                # æ ¹æ®ç¿»è¯‘æ–¹å‘å’Œå†…å®¹è¯­è¨€å†³å®šæ˜¯å¦è·³è¿‡ç¿»è¯‘
                should_skip_translation = False
                if is_target_chinese and is_chinese_content:
                    # è‹±ç¿»ä¸­åœºæ™¯ï¼ŒåŸæ–‡å·²ç»æ˜¯ä¸­æ–‡ï¼Œè·³è¿‡ç¿»è¯‘
                    should_skip_translation = True
                    logger.info(f"è·³è¿‡å·²æ˜¯ä¸­æ–‡çš„æ ‡é¢˜: {text[:30]}...")
                elif is_target_english and is_english_content:
                    # ä¸­ç¿»è‹±åœºæ™¯ï¼ŒåŸæ–‡å·²ç»æ˜¯è‹±æ–‡ï¼Œè·³è¿‡ç¿»è¯‘
                    should_skip_translation = True
                    logger.info(f"è·³è¿‡å·²æ˜¯è‹±æ–‡çš„æ ‡é¢˜: {text[:30]}...")

                translated = cache.get(text)
                if translated is None and not should_skip_translation:
                    translated = _sync_translate_single_text(text, source_language, target_language, custom_translations)
                    cache[text] = translated

                if translated and not should_skip_translation:
                    generator.add_translated_text(translated)
                generator.document.add_paragraph()
                continue

            if btype == 'ul_item':
                generator.add_list_item(text, numbered=False)
                translated = cache.get(text)
                if translated is None:
                    translated = _sync_translate_single_text(text, source_language, target_language, custom_translations)
                    cache[text] = translated
                if translated:
                    generator.add_translated_text(translated)
                generator.document.add_paragraph()
                continue

            if btype == 'ol_item':
                generator.add_list_item(text, numbered=True)
                translated = cache.get(text)
                if translated is None:
                    translated = _sync_translate_single_text(text, source_language, target_language, custom_translations)
                    cache[text] = translated
                if translated:
                    generator.add_translated_text(translated)
                generator.document.add_paragraph()
                continue

            # æ£€æµ‹æ˜¯å¦æ˜¯HTMLè¡¨æ ¼
            if '<table' in text and '</table>' in text:
                # æ·»åŠ åŒè¯­è¡¨æ ¼ï¼ˆåŸæ–‡è¡¨æ ¼å’Œè¯‘æ–‡è¡¨æ ¼ï¼‰
                generator.add_bilingual_table(
                    text,
                    source_language=source_language,
                    target_language=target_language,
                    custom_translations=custom_translations
                )
                continue

            # æ™®é€šæ®µè½
            translated = cache.get(text)
            if translated is None:
                translated = _sync_translate_single_text(text, source_language, target_language, custom_translations)
                cache[text] = translated

            # æ£€æŸ¥æ–‡æœ¬è¯­è¨€ç‰¹æ€§
            is_chinese_content = is_mostly_chinese(text)
            is_target_english = target_language.lower().startswith('en')
            is_target_chinese = target_language.lower().startswith('zh')

            # æ£€æŸ¥æ˜¯å¦æ˜¯è‹±æ–‡å†…å®¹ï¼ˆç®€å•æ£€æµ‹ï¼šåŒ…å«è¾ƒå¤šæ‹‰ä¸å­—æ¯ï¼‰
            def is_mostly_english(s: str) -> bool:
                if not s:
                    return False
                total = len(s)
                if total == 0:
                    return False
                import re
                latin = len(re.findall(r"[a-zA-Z]", s))
                return (latin / total) > 0.5

            is_english_content = is_mostly_english(text)

            # æ ¹æ®ç¿»è¯‘æ–¹å‘å’Œå†…å®¹è¯­è¨€å†³å®šæ˜¯å¦è·³è¿‡ç¿»è¯‘
            should_skip_translation = False
            if is_target_chinese and is_chinese_content:
                # è‹±ç¿»ä¸­åœºæ™¯ï¼ŒåŸæ–‡å·²ç»æ˜¯ä¸­æ–‡ï¼Œè·³è¿‡ç¿»è¯‘
                should_skip_translation = True
                logger.info(f"è·³è¿‡å·²æ˜¯ä¸­æ–‡çš„æ–‡æœ¬: {text[:30]}...")
            elif is_target_english and is_english_content:
                # ä¸­ç¿»è‹±åœºæ™¯ï¼ŒåŸæ–‡å·²ç»æ˜¯è‹±æ–‡ï¼Œè·³è¿‡ç¿»è¯‘
                should_skip_translation = True
                logger.info(f"è·³è¿‡å·²æ˜¯è‹±æ–‡çš„æ–‡æœ¬: {text[:30]}...")

            if translated and not should_skip_translation:
                # æ ¹æ®å†…å®¹å’Œç›®æ ‡è¯­è¨€é€‰æ‹©åˆé€‚çš„æ–‡æ¡£æ ¼å¼
                if is_chinese_content and is_target_english:
                    # ä¸­æ–‡å†…å®¹ç¿»è¯‘ä¸ºè‹±æ–‡ï¼Œç¡®ä¿åŸæ–‡åœ¨ä¸Šï¼Œè¯‘æ–‡åœ¨ä¸‹
                    generator.add_original_text(text)  # æ·»åŠ ä¸­æ–‡åŸæ–‡
                    generator.add_translated_text(translated)  # æ·»åŠ è‹±æ–‡è¯‘æ–‡
                elif is_english_content and is_target_chinese:
                    # è‹±æ–‡å†…å®¹ç¿»è¯‘ä¸ºä¸­æ–‡ï¼Œä½¿ç”¨æ ‡å‡†åŒè¯­å¯¹æ ¼å¼
                    generator.add_bilingual_pair(text, translated)
                else:
                    # å…¶ä»–æƒ…å†µï¼Œä½¿ç”¨æ ‡å‡†åŒè¯­å¯¹æ ¼å¼
                    generator.add_bilingual_pair(text, translated)
            else:
                # è·³è¿‡ç¿»è¯‘æˆ–ç¿»è¯‘å¤±è´¥ï¼Œåªæ·»åŠ åŸæ–‡
                generator.add_original_text(text)
        return generator.save(output_path)
    except Exception as e:
        logger.error(f"translate_markdown_to_bilingual_docå¤„ç†å¤±è´¥: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False
